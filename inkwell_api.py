#!/usr/bin/env python3
"""Inkwell backend API class."""

import webview
import json
import re
import os
import shutil
import base64
import sys
import subprocess
from editor_spoke import get_llm_suggestions
from datetime import datetime
from pathlib import Path
from utils import get_asset_path

# ── Application version ───────────────────────────────────────────────────────
APP_VERSION = '1.0.1'

class Inkwell:
    """Backend API for the writing assistant"""
    
    def __init__(self):
        self.current_project_path = None # We will use this in the next step!
        self.custom_rules = []
        self._window = None
        self._is_maximized = False
        self.load_custom_rules()
        
        # --- NEW: LOAD GLOBAL SETTINGS ---
        self.app_config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'inkwell_app_config.json')
        self.load_global_settings()

    def _load_html(self, filename):
        """Load an HTML file from the ui/ subdirectory."""
        import os
        filepath = get_asset_path(os.path.join('ui', filename))
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()

    def load_global_settings(self):
        """Loads AI keys, themes, holo mode, last project, backup and stats paths."""
        self.stats_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'inkwell_stats.json')
        if os.path.exists(self.app_config_path):
            try:
                with open(self.app_config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.ai_mode           = config.get('ai_mode', 'local')
                    self.ai_api_key        = config.get('ai_api_key', '')
                    self.ai_local_url      = config.get('ai_local_url', 'http://localhost:1234/v1')
                    self.app_theme         = config.get('theme', 'dark')
                    self.holo_mode         = config.get('holo_mode', 'off')
                    self.last_project_path = config.get('last_project_path', '')
                    self.last_project_name = config.get('last_project_name', '')
                    self.backup_folder     = config.get('backup_folder', '')
                    self.backup_keep       = int(config.get('backup_keep', 5))
                    self.last_backup_ts    = config.get('last_backup_ts', '')
                    self.recent_projects   = config.get('recent_projects', [])
                    self.banner_image_path = config.get('banner_image_path', '')
                    self.banner_height     = int(config.get('banner_height', 260))
                    self.banner_scale      = int(config.get('banner_scale',  100))
                    self.banner_pos_y      = int(config.get('banner_pos_y',  0))
                    self.open_satellites    = config.get('open_satellites', [])
                    self.update_check_url   = config.get('update_check_url', '')
                    return
            except Exception as e:
                print(f"Error loading global config: {e}")

        # Defaults if file doesn't exist
        self.ai_mode           = 'local'
        self.ai_api_key        = ''
        self.ai_local_url      = 'http://localhost:1234/v1'
        self.app_theme         = 'dark'
        self.holo_mode         = 'off'
        self.last_project_path = ''
        self.last_project_name = ''
        self.backup_folder     = ''
        self.backup_keep       = 5
        self.last_backup_ts    = ''
        self.recent_projects   = []
        self.banner_image_path = ''
        self.banner_height     = 260
        self.banner_scale      = 100
        self.banner_pos_y      = 0
        self.open_satellites    = []
        self.update_check_url   = ''

    def _save_app_config(self):
        """Write the full app config to disk (shared by all setters)."""
        config = {
            'ai_mode':           getattr(self, 'ai_mode',           'local'),
            'ai_api_key':        getattr(self, 'ai_api_key',        ''),
            'ai_local_url':      getattr(self, 'ai_local_url',      'http://localhost:1234/v1'),
            'theme':             getattr(self, 'app_theme',         'dark'),
            'holo_mode':         getattr(self, 'holo_mode',         'off'),
            'last_project_path': getattr(self, 'last_project_path', ''),
            'last_project_name': getattr(self, 'last_project_name', ''),
            'backup_folder':     getattr(self, 'backup_folder',     ''),
            'backup_keep':       getattr(self, 'backup_keep',       5),
            'last_backup_ts':    getattr(self, 'last_backup_ts',    ''),
            'recent_projects':   getattr(self, 'recent_projects',   []),
            'banner_image_path': getattr(self, 'banner_image_path', ''),
            'banner_height':     getattr(self, 'banner_height',     260),
            'banner_scale':      getattr(self, 'banner_scale',      100),
            'banner_pos_y':      getattr(self, 'banner_pos_y',      0),
            'open_satellites':   getattr(self, 'open_satellites',   []),
            'update_check_url':  getattr(self, 'update_check_url',  ''),
        }
        with open(self.app_config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=4)

    # ── BACKUP ──────────────────────────────────────────────────────────────

    def get_backup_settings(self):
        """Return current backup configuration so the UI can display it."""
        return {
            'folder':      getattr(self, 'backup_folder',  ''),
            'keep':        getattr(self, 'backup_keep',    5),
            'last_backup': getattr(self, 'last_backup_ts', ''),
        }

    def pick_backup_folder(self):
        """Open a native folder-picker so the user can choose the backup destination.
        Persists the chosen path and returns it (or None on cancel)."""
        if not self._window:
            return None
        result = self._window.create_file_dialog(webview.FileDialog.FOLDER, allow_multiple=False)
        if result and len(result) > 0:
            folder = result[0]
            self.backup_folder = folder
            self._save_app_config()
            return folder
        return None

    def save_backup(self, keep=None):
        """Copy the whole project folder into a timestamped sub-directory inside
        backup_folder and prune old copies to stay within the keep limit.

        Returns:
            { success: bool, path: str, timestamp: str, error: str }
        """
        if not getattr(self, 'current_project_path', None):
            return {'success': False, 'error': 'No project open'}

        folder = getattr(self, 'backup_folder', '').strip()
        if not folder or not os.path.isdir(folder):
            return {'success': False, 'error': 'Backup folder not set or does not exist'}

        if keep is not None:
            try:
                self.backup_keep = max(1, int(keep))
                self._save_app_config()
            except (ValueError, TypeError):
                pass

        keep_n = max(1, getattr(self, 'backup_keep', 5))

        # Derive a safe project name from the folder name
        proj_name = os.path.basename(self.current_project_path.rstrip('/\\'))
        ts        = datetime.now().strftime('%Y%m%d_%H%M%S')
        dest_name = f'{proj_name}_{ts}'
        dest_path = os.path.join(folder, dest_name)

        try:
            # Copy the project tree; skip the checkpoints/ subfolder (transient data)
            def _ignore(src, names):
                return ['checkpoints'] if os.path.basename(src) == os.path.basename(self.current_project_path) else []

            shutil.copytree(self.current_project_path, dest_path, ignore=_ignore)

            # Prune oldest backups for this project so we keep at most keep_n copies
            prefix = proj_name + '_'
            all_backups = sorted(
                [d for d in os.listdir(folder)
                 if d.startswith(prefix) and os.path.isdir(os.path.join(folder, d))],
            )
            while len(all_backups) > keep_n:
                oldest = os.path.join(folder, all_backups.pop(0))
                shutil.rmtree(oldest, ignore_errors=True)

            # Persist timestamp
            self.last_backup_ts = datetime.now().isoformat()
            self._save_app_config()

            human_ts = datetime.now().strftime('%b %d, %Y at %I:%M %p')
            print(f'[BACKUP] Saved to: {dest_path}  (keeping last {keep_n})')
            return {'success': True, 'path': dest_path, 'timestamp': human_ts}

        except Exception as e:
            print(f'[BACKUP] Error: {e}')
            return {'success': False, 'error': str(e)}

    def save_backup_settings(self, folder, keep):
        """Persist backup folder + keep count without triggering an immediate backup."""
        self.backup_folder = str(folder).strip()
        try:
            self.backup_keep = max(1, int(keep))
        except (ValueError, TypeError):
            self.backup_keep = 5
        self._save_app_config()
        return {'success': True}

    # ── WRITING STATS ────────────────────────────────────────────────────────

    def _load_stats(self):
        """Return the stats dict, or a clean default if the file is missing/corrupt."""
        sp = getattr(self, 'stats_path', None)
        if sp and os.path.exists(sp):
            try:
                with open(sp, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                pass
        return {
            'daily_counts':    {},
            'hourly_counts':   {},
            'personal_bests':  {'most_words_day': 0, 'longest_session_words': 0, 'fastest_wpm': 0},
            'current_streak':  0,
            'streak_last_date': '',
            'longest_streak':  0,
        }

    def _save_stats(self, stats):
        sp = getattr(self, 'stats_path', None)
        if not sp:
            return
        with open(sp, 'w', encoding='utf-8') as f:
            json.dump(stats, f, indent=2)

    def get_writing_stats(self):
        """Return the full stats object for the writing stats dashboard."""
        from datetime import date, timedelta
        stats = self._load_stats()
        today = date.today().isoformat()
        # Attach today's count for convenience
        stats['today_count'] = stats['daily_counts'].get(today, 0)
        # Build last 7 days array for the bar chart
        week = []
        for i in range(6, -1, -1):
            d = (date.today() - timedelta(days=i)).isoformat()
            week.append({'date': d, 'count': stats['daily_counts'].get(d, 0)})
        stats['week'] = week
        # Pull word goal from the current project settings if available
        word_goal = 0
        try:
            if getattr(self, 'current_project_path', '') and os.path.exists(self.current_project_path):
                with open(self.current_project_path, 'r', encoding='utf-8') as f:
                    proj = json.load(f)
                word_goal = int((proj.get('settings') or {}).get('wordGoal') or 0)
        except Exception:
            pass
        stats['word_goal'] = word_goal
        return stats

    def record_session_words(self, delta, session_total, wpm):
        """
        Record words written. Called periodically from JS.
        delta        — words added since last call
        session_total — total words typed this session (for PB comparison)
        wpm           — current rolling WPM
        Returns the updated stats dict plus any new_pbs list.
        """
        from datetime import date, timedelta
        stats = self._load_stats()
        today     = date.today().isoformat()
        yesterday = (date.today() - timedelta(days=1)).isoformat()
        hour      = str(datetime.now().hour)

        delta         = max(0, int(delta))
        session_total = max(0, int(session_total))
        wpm           = max(0, int(wpm))

        # ── Daily count ──────────────────────────────────────────────────────
        if delta > 0:
            stats['daily_counts'][today] = stats['daily_counts'].get(today, 0) + delta
            stats['hourly_counts'][hour] = stats['hourly_counts'].get(hour, 0) + delta

        # ── Streak ───────────────────────────────────────────────────────────
        last = stats.get('streak_last_date', '')
        if last != today and delta > 0:
            if last == yesterday:
                stats['current_streak'] = stats.get('current_streak', 0) + 1
            elif last != '':
                stats['current_streak'] = 1   # streak broken
            else:
                stats['current_streak'] = 1   # first ever day
            stats['streak_last_date'] = today
        stats['longest_streak'] = max(stats.get('longest_streak', 0),
                                      stats.get('current_streak', 0))

        # ── Personal bests ───────────────────────────────────────────────────
        pb = stats.setdefault('personal_bests',
                              {'most_words_day': 0, 'longest_session_words': 0, 'fastest_wpm': 0})
        new_pbs = []

        day_total = stats['daily_counts'].get(today, 0)
        if day_total > pb.get('most_words_day', 0):
            pb['most_words_day'] = day_total
            new_pbs.append({'type': 'most_words_day', 'value': day_total})

        if session_total > pb.get('longest_session_words', 0):
            pb['longest_session_words'] = session_total
            new_pbs.append({'type': 'longest_session_words', 'value': session_total})

        if wpm > pb.get('fastest_wpm', 0) and wpm > 10:
            pb['fastest_wpm'] = wpm
            new_pbs.append({'type': 'fastest_wpm', 'value': wpm})

        self._save_stats(stats)

        result = dict(stats)
        result['today_count'] = day_total
        if new_pbs:
            result['new_pbs'] = new_pbs
        return result

    # ── HOLO ────────────────────────────────────────────────────────────────

    def save_holo_mode(self, mode):
        """Persist the active hologram tier so splash + next launch can restore it."""
        valid = ('off', 'holo-lite', 'holo', 'holo-max')
        self.holo_mode = mode if mode in valid else 'off'
        self._save_app_config()
        print(f"[HOLO] Mode saved: {self.holo_mode}")
        return True

    def get_holo_mode(self):
        """Return the last-saved hologram tier."""
        return getattr(self, 'holo_mode', 'off')

    def get_last_project(self):
        """Return the last opened project path and display name for the welcome screen."""
        return {
            'path': getattr(self, 'last_project_path', ''),
            'name': getattr(self, 'last_project_name', ''),
        }

    def open_last_project(self):
        """Open the most recently used project directly, bypassing the file picker."""
        path = getattr(self, 'last_project_path', '')
        if not path or not os.path.exists(path):
            return {"success": False, "error": "Last project not found or has been moved."}
        return self.open_project(path)

    # ── RECENT PROJECTS ─────────────────────────────────────────────────────

    def _push_recent_project(self, path, name):
        """Insert/update a project entry at the front of the recent-projects list (max 8)."""
        rp = getattr(self, 'recent_projects', [])
        # Remove any existing entry for this path
        rp = [e for e in rp if e.get('path') != path]
        rp.insert(0, {'path': path, 'name': name})
        self.recent_projects = rp[:8]

    def get_recent_projects(self):
        """Return enriched recent-project cards for the welcome screen grid."""
        import glob as _glob
        rp = getattr(self, 'recent_projects', [])
        cards = []
        for entry in rp:
            meta_file = entry.get('path', '')
            name      = entry.get('name', '')
            if not meta_file or not os.path.exists(meta_file):
                continue
            # Quick meta read for word count + chapter count
            try:
                with open(meta_file, 'r', encoding='utf-8') as f:
                    meta = json.load(f)
                chapters = meta.get('chapters', [])
                wc = 0
                for chap in chapters:
                    # Content is NOT loaded here (stored separately) — use chapter dir
                    chap_id = chap.get('id', '')
                    if chap_id:
                        chap_path = os.path.join(os.path.dirname(meta_file), 'chapters', chap_id + '.html')
                        if os.path.exists(chap_path):
                            with open(chap_path, 'r', encoding='utf-8') as cf:
                                raw = re.sub(r'<[^>]+>', ' ', cf.read())
                                wc += len(raw.split())
                mod_ts   = os.path.getmtime(meta_file)
                modified = datetime.fromtimestamp(mod_ts).strftime('%b %d, %Y')
                cards.append({
                    'name':      meta.get('title', name),
                    'meta_file': meta_file,
                    'words':     wc,
                    'chapters':  len(chapters),
                    'modified':  modified,
                })
            except Exception:
                continue
        return cards

    def pick_and_scan_projects(self):
        """Open a native folder picker then scan it for .inkwell projects."""
        if not self._window:
            return {'success': False, 'error': 'No window'}
        result = self._window.create_file_dialog(webview.FileDialog.FOLDER, allow_multiple=False)
        if not result or not len(result):
            return {'success': False, 'error': 'Cancelled'}
        res = self.scan_projects_folder(result[0])
        if res.get('success'):
            # Map to the same shape used by get_recent_projects so JS is consistent
            for c in res['cards']:
                c['name'] = c.pop('title', c.get('name', ''))
        return res

    # ── BANNER IMAGE SETTINGS ──────────────────────────────────────────────────

    def _image_to_data_uri(self, path):
        """Read an image file and return a data: URI string, or None on error."""
        try:
            ext  = os.path.splitext(path)[1].lower().lstrip('.')
            mime = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                    'png': 'image/png',  'gif': 'image/gif',
                    'webp': 'image/webp', 'bmp': 'image/bmp'}.get(ext, 'image/jpeg')
            import base64 as _b64
            with open(path, 'rb') as f:
                return 'data:' + mime + ';base64,' + _b64.b64encode(f.read()).decode('ascii')
        except Exception:
            return None

    def get_banner_settings(self):
        """Return stored banner display settings and the data URI for any custom image."""
        path     = getattr(self, 'banner_image_path', '')
        data_uri = self._image_to_data_uri(path) if path and os.path.exists(path) else None
        return {
            'image_path':    path,
            'image_data_uri': data_uri,
            'height': getattr(self, 'banner_height', 260),
            'scale':  getattr(self, 'banner_scale',  100),
            'pos_y':  getattr(self, 'banner_pos_y',  0),
        }

    def pick_banner_image(self):
        """Open a file picker for images and return the path + data URI."""
        if not self._window:
            return None
        result = self._window.create_file_dialog(
            webview.FileDialog.OPEN,
            allow_multiple=False,
            file_types=('Image Files (*.png;*.jpg;*.jpeg;*.gif;*.webp;*.bmp)', 'All Files (*.*)')
        )
        if not result or not len(result):
            return None
        path     = result[0]
        data_uri = self._image_to_data_uri(path)
        if not data_uri:
            return {'error': 'Could not read image file'}
        return {'path': path, 'data_uri': data_uri}

    def save_banner_settings(self, image_path, height, scale, pos_y):
        """Persist banner display preferences to the app config."""
        self.banner_image_path = str(image_path) if image_path else ''
        try: self.banner_height = max(150, min(500, int(height)))
        except Exception: self.banner_height = 260
        try: self.banner_scale  = max(100, min(300, int(scale)))
        except Exception: self.banner_scale  = 100
        try: self.banner_pos_y  = max(0,   min(100, int(pos_y)))
        except Exception: self.banner_pos_y  = 0
        self._save_app_config()
        return True

    # ── INLINE GHOST-TEXT ───────────────────────────────────────────────────
    def get_inline_suggestion(self, context):
        """Return a short prose continuation for the ghost-text feature.

        Calls the same OpenAI-compatible endpoint used by the rest of Inkwell.
        Returns {'suggestion': str} on success or {'suggestion': '', 'error': str}.
        Uses a hard 8-second timeout and max_tokens=60 to stay snappy.
        """
        import urllib.request, urllib.error

        if not context or len(context.strip()) < 40:
            return {'suggestion': ''}

        context = context[-500:]   # cap context window for speed

        system_prompt = (
            "You are a ghost-text autocomplete engine inside a fiction writing app. "
            "Continue the author's text naturally with at most one sentence (15 words max). "
            "Output ONLY the continuation — no quotes, no explanation, no prefix. "
            "Match the author's voice, tense, and point of view exactly."
        )

        mode    = getattr(self, 'ai_mode',      'local')
        api_key = getattr(self, 'ai_api_key',   '').strip()
        url_box = getattr(self, 'ai_local_url', 'http://localhost:1234/v1').strip().rstrip('/')

        if mode == 'local':
            url = url_box if '/chat/completions' in url_box else url_box + '/v1/chat/completions'
            if not url_box:
                url = 'http://localhost:1234/v1/chat/completions'
        else:
            url = 'https://api.groq.com/openai/v1/chat/completions'

        payload = {
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user',   'content': context}
            ],
            'temperature': 0.72,
            'max_tokens':  60,
        }
        if mode != 'local':
            payload['model'] = 'llama-3.3-70b-versatile'

        try:
            req = urllib.request.Request(url, method='POST')
            req.add_header('Content-Type', 'application/json')
            req.add_header('User-Agent', 'Inkwell/1.0')
            if mode != 'local' and api_key:
                req.add_header('Authorization', f'Bearer {api_key}')

            with urllib.request.urlopen(req, data=json.dumps(payload).encode(), timeout=8) as resp:
                body       = json.loads(resp.read().decode())
                suggestion = body['choices'][0]['message']['content'].strip().strip('"\'')
                # If model echoed the context back, return empty
                if suggestion.lower().startswith(context[-30:].lower().strip()):
                    return {'suggestion': ''}
                return {'suggestion': suggestion}
        except Exception as e:
            return {'suggestion': '', 'error': str(e)}

    # ── SLASH COMMANDS ─────────────────────────────────────────────────────
    def _slash_call(self, system_prompt, user_text, max_tokens=250):
        """Shared HTTP helper for all slash-command API calls."""
        import urllib.request

        mode    = getattr(self, 'ai_mode',      'local')
        api_key = getattr(self, 'ai_api_key',   '').strip()
        url_box = getattr(self, 'ai_local_url', 'http://localhost:1234/v1').strip().rstrip('/')

        if mode == 'local':
            url = url_box if '/chat/completions' in url_box else url_box + '/v1/chat/completions'
            if not url_box:
                url = 'http://localhost:1234/v1/chat/completions'
        else:
            url = 'https://api.groq.com/openai/v1/chat/completions'

        payload = {
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user',   'content': user_text},
            ],
            'temperature': 0.80,
            'max_tokens':  max_tokens,
        }
        if mode != 'local':
            payload['model'] = 'llama-3.3-70b-versatile'

        req = urllib.request.Request(url, method='POST')
        req.add_header('Content-Type', 'application/json')
        req.add_header('User-Agent', 'Inkwell/1.0')
        if mode != 'local' and api_key:
            req.add_header('Authorization', f'Bearer {api_key}')

        with urllib.request.urlopen(req, data=json.dumps(payload).encode(), timeout=20) as resp:
            body = json.loads(resp.read().decode())
            return body['choices'][0]['message']['content'].strip()

    def slash_continue(self, paragraph, context=''):
        """Write the next 2-3 sentences continuing the story after the paragraph."""
        system = (
            "You are a fiction writing assistant. "
            "Continue the story naturally from where the excerpt ends. "
            "Write 2–3 sentences that flow seamlessly from the last line. "
            "Output ONLY the continuation text — no labels, no explanation."
        )
        user = (f"Context:\n{context}\n\n" if context else '') + \
               f"Continue from:\n{paragraph}"
        try:
            return {'result': self._slash_call(system, user, max_tokens=150)}
        except Exception as e:
            return {'result': '', 'error': str(e)}

    def slash_rephrase(self, paragraph, context=''):
        """Rewrite the paragraph with fresher prose, same meaning and voice."""
        system = (
            "You are a fiction editor. "
            "Rewrite the paragraph with the same meaning but fresher, more vivid prose. "
            "Preserve the author's voice, tense, and point of view. "
            "Output ONLY the rewritten paragraph — no explanation, no preamble."
        )
        try:
            return {'result': self._slash_call(system, paragraph, max_tokens=300)}
        except Exception as e:
            return {'result': '', 'error': str(e)}

    def slash_suggest(self, paragraph, context=''):
        """Return a prose-quality-improved version of the paragraph."""
        system = (
            "You are a fiction editor focused on prose quality. "
            "Improve the paragraph: sharpen imagery, cut weak words, "
            "vary sentence rhythm, and strengthen verbs. "
            "Preserve the author's voice and meaning. "
            "Output ONLY the improved paragraph — no explanation."
        )
        try:
            return {'result': self._slash_call(system, paragraph, max_tokens=300)}
        except Exception as e:
            return {'result': '', 'error': str(e)}

    # ── AI CONSISTENCY CHECKER ─────────────────────────────────────────────
    def check_consistency(self):
        """Batch-scan all chapters for continuity errors, timeline gaps,
        and attribute contradictions.  Returns {success, issues} where each
        issue is {type, severity, issue, chapters, detail}."""
        import re

        if not getattr(self, 'current_project_path', None):
            return {'success': False, 'error': 'No project is open.'}

        # ── 1. Read all chapters ──────────────────────────────────────────
        try:
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                db = json.load(f)
        except Exception as e:
            return {'success': False, 'error': f'Could not read project: {e}'}

        chapters = db.get('chapters', [])
        if not chapters:
            return {'success': True, 'issues': []}

        chap_dir = os.path.join(self.current_project_path, 'chapters')

        # Strip HTML tags helper
        def _strip(html):
            return re.sub(r'<[^>]+>', ' ', html or '').strip()

        # Build a condensed manuscript: chapter title + first ~300 words each
        manuscript_parts = []
        for i, chap in enumerate(chapters):
            chap_file = os.path.join(chap_dir, f"{chap['id']}.html")
            content = ''
            if os.path.exists(chap_file):
                with open(chap_file, 'r', encoding='utf-8') as cf:
                    content = _strip(cf.read())
            # Take first 300 words
            words = content.split()[:300]
            excerpt = ' '.join(words)
            title = chap.get('title', f'Chapter {i+1}')
            manuscript_parts.append(f'[Chapter {i+1}: "{title}"]\n{excerpt}')

        manuscript = '\n\n'.join(manuscript_parts)

        # Also inject character profiles if available
        char_path = os.path.join(self.current_project_path, 'characters.json')
        char_profiles = ''
        if os.path.exists(char_path):
            try:
                with open(char_path, 'r', encoding='utf-8') as f:
                    chars = json.load(f)
                char_lines = []
                for c in chars[:20]:  # cap at 20 characters
                    name = c.get('name', '')
                    glance = _strip(c.get('glance', ''))
                    if name:
                        char_lines.append(f"- {name}: {glance[:120]}")
                if char_lines:
                    char_profiles = '\n'.join(char_lines)
            except Exception:
                pass

        # ── 2. Build the prompt ───────────────────────────────────────────
        system_prompt = (
            "You are a meticulous fiction continuity editor. "
            "Given a condensed manuscript (chapter excerpts) and optional character profiles, "
            "identify potential continuity problems. "
            "Look for: character attribute contradictions (eye colour, skills, relationships), "
            "timeline or chronology gaps, location contradictions, "
            "plot-logic inconsistencies, and object/item continuity breaks. "
            "Respond ONLY with a valid JSON array. "
            "Each element must have exactly these keys: "
            '{"type": string, "severity": "high"|"medium"|"low", '
            '"issue": string (one sentence), '
            '"chapters": [list of chapter numbers as integers], '
            '"detail": string (1-2 sentences of detail)}. '
            "If there are no issues, return an empty array []. "
            "Do NOT include commentary outside the JSON."
        )

        user_text = ''
        if char_profiles:
            user_text += f"=== CHARACTER PROFILES ===\n{char_profiles}\n\n"
        user_text += f"=== MANUSCRIPT EXCERPTS ===\n{manuscript}"

        # ── 3. Call the AI ────────────────────────────────────────────────
        try:
            raw = self._slash_call(system_prompt, user_text, max_tokens=1200)
        except Exception as e:
            return {'success': False, 'error': f'AI request failed: {e}'}

        # ── 4. Parse JSON robustly ────────────────────────────────────────
        import ast
        json_match = re.search(r'\[.*\]', raw, re.DOTALL)
        if not json_match:
            return {'success': True, 'issues': []}
        clean = json_match.group(0)
        try:
            issues = json.loads(clean)
        except json.JSONDecodeError:
            try:
                issues = ast.literal_eval(clean)
            except Exception as e:
                return {'success': False, 'error': f'Parse error: {e}'}

        if not isinstance(issues, list):
            issues = []

        return {'success': True, 'issues': issues}

    # ==========================================
    # PROJECT BUNDLE ENGINE (V1.0 Architecture)
    # ==========================================
    def create_new_project(self, folder_path, project_name):
        """Creates the professional directory structure for a new novel."""
        self.current_project_path = folder_path
        
        # Extract "My_Masterpiece" from "My_Masterpiece.inkwell"
        base_filename = os.path.basename(folder_path).replace(".inkwell", "")
        
        # Save the exact path to this new dynamically named file!
        self.current_meta_path = os.path.join(folder_path, f"{base_filename}.json")
        
        folders_to_create = ['chapters', 'snapshots', 'exports', 'assets']
        for folder in folders_to_create:
            os.makedirs(os.path.join(folder_path, folder), exist_ok=True)
            
        project_metadata = {
            "title": project_name,
            "created_at": str(datetime.now()),
            "chapters": [],       
            "characters": [],     
            "world_notes": {}     
        }
        
        # Save using the new dynamic path
        with open(self.current_meta_path, 'w', encoding='utf-8') as f:
            json.dump(project_metadata, f, indent=4)

        # Persist as the "last opened" project so the welcome screen can offer a quick resume
        self.last_project_path = self.current_meta_path
        self.last_project_name = project_name
        self._push_recent_project(self.current_meta_path, project_name)
        self._save_app_config()

        return {"success": True, "data": project_metadata}

    def open_project(self, file_path):
        """Loads a project bundle from the dynamically named file."""
        import os, json
        
        # Store BOTH the file path and the parent folder path in memory
        self.current_meta_path = file_path
        self.current_project_path = os.path.dirname(file_path)
        
        if not os.path.exists(self.current_meta_path):
            return {"success": False, "error": "Not a valid Inkwell project file."}
            
        with open(self.current_meta_path, 'r', encoding='utf-8') as f:
            db = json.load(f)
            
        # 1. Load Chapters
        chap_dir = os.path.join(self.current_project_path, "chapters")
        for chap in db.get('chapters', []):
            chap_file = os.path.join(chap_dir, f"{chap['id']}.html")
            if os.path.exists(chap_file):
                with open(chap_file, 'r', encoding='utf-8') as cf:
                    chap['content'] = cf.read()
            else:
                chap['content'] = "" 
                
        # 2. Load Satellites
        char_path = os.path.join(self.current_project_path, "characters.json")
        if os.path.exists(char_path):
            with open(char_path, 'r', encoding='utf-8') as f:
                db['characters'] = json.load(f)
                
        notes_path = os.path.join(self.current_project_path, "world_notes.json")
        if os.path.exists(notes_path):
            with open(notes_path, 'r', encoding='utf-8') as f:
                db['world_notes'] = json.load(f)

        # Persist as the "last opened" project
        self.last_project_path = file_path
        self.last_project_name = db.get('title', '')
        self._push_recent_project(file_path, db.get('title', ''))
        self._save_app_config()

        result = {"success": True, "data": db}

        # Check for a crash-recovery checkpoint that is newer than the saved meta
        checkpoint_info = self.get_latest_checkpoint()
        if checkpoint_info:
            result["checkpoint"] = checkpoint_info

        return result

    def save_chapter_file(self, chapter_id, content):
        """Saves a single chapter to its own HTML file."""
        if not getattr(self, 'current_project_path', None): 
            return {"success": False}
            
        try:
            chap_dir = os.path.join(self.current_project_path, "chapters")
            os.makedirs(chap_dir, exist_ok=True)
            
            # Ensure it's a string, write empty if None
            safe_content = str(content) if content else ""
            
            file_path = os.path.join(chap_dir, f"{chapter_id}.html")
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(safe_content)
                
            return {"success": True}
        except Exception as e:
            print(f"Chapter file save error: {e}")
            return {"success": False}
    
    def delete_chapter_file(self, chapter_id):
        """Removes the standalone HTML file when a chapter is deleted from the binder."""
        if not self.current_project_path: 
            return {"success": False}
            
        file_path = os.path.join(self.current_project_path, "chapters", f"{chapter_id}.html")
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                return {"success": True}
            except Exception as e:
                print(f"Could not delete chapter file: {e}")
                
        return {"success": False}

    def save_characters(self, char_data):
        """Saves ONLY the Character Bible to its dedicated file in the project bundle."""
        if not self.current_project_path:
            return {"success": False, "error": "No project open."}

        char_path = os.path.join(self.current_project_path, "characters.json")
        with open(char_path, 'w', encoding='utf-8') as f:
            json.dump(char_data, f, indent=4)
        # Refresh lore hover index AND dock Cast panel in main window
        try:
            import webview as _wv
            for w in _wv.windows:
                if w.title == 'Inkwell':
                    w.evaluate_js(
                        "if(window.refreshLoreIndex)window.refreshLoreIndex();"
                        "if(window.refreshDockLists)window.refreshDockLists('characters');"
                    )
                    break
        except Exception:
            pass
        return {"success": True}

    def save_world_notes(self, notes_data):
        """Saves ONLY the World Notes to its dedicated file in the project bundle."""
        if not self.current_project_path:
            return {"success": False, "error": "No project open."}

        notes_path = os.path.join(self.current_project_path, "world_notes.json")
        with open(notes_path, 'w', encoding='utf-8') as f:
            json.dump(notes_data, f, indent=4)
        # Refresh lore hover index AND dock World panel in main window
        try:
            import webview as _wv
            for w in _wv.windows:
                if w.title == 'Inkwell':
                    w.evaluate_js(
                        "if(window.refreshLoreIndex)window.refreshLoreIndex();"
                        "if(window.refreshDockLists)window.refreshDockLists('world');"
                    )
                    break
        except Exception:
            pass
        return {"success": True}

    def get_characters(self):
        """Return the full character list from disk (used by the main window dock refresh)."""
        if not self.current_project_path:
            return []
        char_path = os.path.join(self.current_project_path, 'characters.json')
        if os.path.exists(char_path):
            try:
                with open(char_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return data if isinstance(data, list) else []
            except Exception:
                pass
        return []

    def get_world_notes(self):
        """Return the full world-notes dict from disk (used by the main window dock refresh)."""
        if not self.current_project_path:
            return {}
        notes_path = os.path.join(self.current_project_path, 'world_notes.json')
        if os.path.exists(notes_path):
            try:
                with open(notes_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                return data if isinstance(data, dict) else {}
            except Exception:
                pass
        return {}

    def get_lore_index(self):
        """Return character names and world-note titles for the lore hover-card system."""
        result = {'characters': [], 'notes': []}
        if not self.current_project_path:
            return result

        # Characters
        char_path = os.path.join(self.current_project_path, 'characters.json')
        if os.path.exists(char_path):
            try:
                with open(char_path, 'r', encoding='utf-8') as f:
                    chars = json.load(f)
                if isinstance(chars, list):
                    for c in chars:
                        result['characters'].append({
                            'id':     c.get('id', ''),
                            'name':   c.get('name', ''),
                            'role':   c.get('role', ''),
                            'glance': (c.get('glance') or '')[:200],
                            'avatar': c.get('avatar', ''),
                        })
            except Exception:
                pass

        # World notes
        notes_path = os.path.join(self.current_project_path, 'world_notes.json')
        if os.path.exists(notes_path):
            try:
                with open(notes_path, 'r', encoding='utf-8') as f:
                    world_notes = json.load(f)
                if isinstance(world_notes, dict):
                    for category, notes in world_notes.items():
                        for note in (notes if isinstance(notes, list) else []):
                            title = note.get('title', '')
                            if not title:
                                continue
                            result['notes'].append({
                                'id':       note.get('id', title),
                                'title':    title,
                                'category': category,
                                'content':  (note.get('content') or '')[:200],
                            })
            except Exception:
                pass

        return result

    # ==========================================
    # CRASH RECOVERY — CHECKPOINT ENGINE
    # ==========================================

    def save_checkpoint(self, chapter_contents_json):
        """Write a rolling crash-recovery checkpoint to the project's checkpoints/ folder.
        Called from JS every 60 s with the full in-memory chapter content so the
        checkpoint is always newer than the debounced chapter files on disk.
        Keeps only the 3 most recent checkpoints to avoid disk bloat.
        """
        if not getattr(self, 'current_project_path', None):
            return {"success": False}
        if not getattr(self, 'current_meta_path', None):
            return {"success": False}

        try:
            import copy

            chk_dir = os.path.join(self.current_project_path, 'checkpoints')
            os.makedirs(chk_dir, exist_ok=True)

            # Read current metadata (structure, chapter order, settings)
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                meta = json.load(f)

            # Merge live content from the JS frontend
            chapter_contents = json.loads(chapter_contents_json)
            content_map = {c['id']: c.get('content', '') for c in chapter_contents}

            checkpoint = copy.deepcopy(meta)
            checkpoint['_checkpoint_ts'] = datetime.now().isoformat()
            for chap in checkpoint.get('chapters', []):
                chap['content'] = content_map.get(chap['id'], '')

            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'checkpoint_{ts}.json'
            with open(os.path.join(chk_dir, filename), 'w', encoding='utf-8') as f:
                json.dump(checkpoint, f, indent=2)

            # Prune: keep only the 3 most recent
            all_chk = sorted(
                f for f in os.listdir(chk_dir)
                if f.startswith('checkpoint_') and f.endswith('.json')
            )
            for old in all_chk[:-3]:
                try:
                    os.remove(os.path.join(chk_dir, old))
                except Exception:
                    pass

            return {"success": True, "filename": filename}
        except Exception as e:
            print(f"[CHECKPOINT] save_checkpoint error: {e}")
            return {"success": False, "error": str(e)}

    def get_latest_checkpoint(self):
        """Return info about the most recent checkpoint if it is newer than the
        saved project meta — meaning the last session may have had unsaved work.
        Returns None if no checkpoint exists or the checkpoint is stale.
        """
        if not getattr(self, 'current_project_path', None):
            return None
        if not getattr(self, 'current_meta_path', None):
            return None

        chk_dir = os.path.join(self.current_project_path, 'checkpoints')
        if not os.path.exists(chk_dir):
            return None

        try:
            all_chk = sorted(
                f for f in os.listdir(chk_dir)
                if f.startswith('checkpoint_') and f.endswith('.json')
            )
            if not all_chk:
                return None

            latest_file = os.path.join(chk_dir, all_chk[-1])
            meta_mtime  = os.path.getmtime(self.current_meta_path)
            chk_mtime   = os.path.getmtime(latest_file)

            if chk_mtime <= meta_mtime:
                return None  # checkpoint is older than saved state — nothing to recover

            with open(latest_file, 'r', encoding='utf-8') as f:
                data = json.load(f)

            return {
                "filename":      all_chk[-1],
                "timestamp":     data.get('_checkpoint_ts', all_chk[-1]),
                "chapter_count": len(data.get('chapters', [])),
            }
        except Exception as e:
            print(f"[CHECKPOINT] get_latest_checkpoint error: {e}")
            return None

    def restore_from_checkpoint(self, filename):
        """Load a named checkpoint, write its contents back to the project files,
        and return the restored project data in the same shape as open_project.
        """
        if not getattr(self, 'current_project_path', None):
            return {"success": False, "error": "No project open."}

        try:
            import copy

            chk_file = os.path.join(self.current_project_path, 'checkpoints', filename)
            if not os.path.exists(chk_file):
                return {"success": False, "error": "Checkpoint file not found."}

            with open(chk_file, 'r', encoding='utf-8') as f:
                checkpoint = json.load(f)

            # Remove internal marker before returning to JS
            checkpoint.pop('_checkpoint_ts', None)

            # Write restored meta (without content so the file stays lightweight)
            meta = copy.deepcopy(checkpoint)
            for chap in meta.get('chapters', []):
                chap.pop('content', None)
            with open(self.current_meta_path, 'w', encoding='utf-8') as f:
                json.dump(meta, f, indent=4)

            # Write restored chapter HTML files
            chap_dir = os.path.join(self.current_project_path, 'chapters')
            os.makedirs(chap_dir, exist_ok=True)
            for chap in checkpoint.get('chapters', []):
                chap_file = os.path.join(chap_dir, f"{chap['id']}.html")
                with open(chap_file, 'w', encoding='utf-8') as f:
                    f.write(chap.get('content', ''))

            return {"success": True, "data": checkpoint}
        except Exception as e:
            print(f"[CHECKPOINT] restore_from_checkpoint error: {e}")
            return {"success": False, "error": str(e)}

    def save_project_meta(self, db_data):
        """Saves ONLY the lightweight metadata to the JSON."""
        if not getattr(self, 'current_meta_path', None):
            return {"success": False}
            
        try:
            import copy
            safe_db = copy.deepcopy(db_data) 
            for chap in safe_db.get('chapters', []):
                if 'content' in chap:
                    del chap['content'] 
                    
            # Use the dynamic file path!
            with open(self.current_meta_path, 'w', encoding='utf-8') as f:
                json.dump(safe_db, f, indent=4)
                
            return {"success": True}
        except Exception as e:
            print(f"Meta save error: {e}")
            return {"success": False}
    
    def pick_new_project_location(self):
        """Native dialog to select where to save a new project bundle."""
        import webview
        if not self._window: return None
        # CHANGED: Updated to the modern pywebview syntax
        result = self._window.create_file_dialog(webview.FileDialog.FOLDER, allow_multiple=False)
        if result and len(result) > 0:
            return result[0]
        return None

    def pick_open_project_location(self):
        """Native dialog to open an existing project by selecting its anchor file."""
        import webview
        if not self._window: return None
        
        # FIX: Restored the proper strict formatting for pywebview
        file_types = ('Inkwell Project (*.json)', 'All Files (*.*)')
        
        result = self._window.create_file_dialog(
            webview.FileDialog.OPEN, 
            allow_multiple=False,
            file_types=file_types
        )
        
        if result and len(result) > 0:
            # We NO LONGER strip the filename! Just return the exact file path.
            return result[0] 
        return None
    
    def get_project_lore(self):
        """Helper for satellite windows to fetch lore without a full reload."""
        if not getattr(self, 'current_meta_path', None): return None
        
        # Use the dynamic file path!
        with open(self.current_meta_path, 'r', encoding='utf-8') as f:
            db = json.load(f)

        # Load the external files to sync them
        char_path = os.path.join(self.current_project_path, "characters.json")
        if os.path.exists(char_path):
            with open(char_path, 'r', encoding='utf-8') as f:
                db['characters'] = json.load(f)

        notes_path = os.path.join(self.current_project_path, "world_notes.json")
        if os.path.exists(notes_path):
            with open(notes_path, 'r', encoding='utf-8') as f:
                db['world_notes'] = json.load(f)
                
        return db
    
    # --- MEDIA MANAGEMENT ENGINE ---
    def _get_media_dir(self):
        """Creates and returns the path to the project-specific assets folder"""
        
        # 1. Fallback if somehow triggered without a project open
        if not getattr(self, 'current_project_path', None):
            print("[WARNING] Saving media without an active project.")
            base_dir = os.path.dirname(os.path.abspath(__file__))
            fallback_dir = os.path.join(base_dir, 'Inkwell_Orphaned_Media')
            os.makedirs(fallback_dir, exist_ok=True)
            return fallback_dir
            
        # 2. Route all images to the specific novel's internal folder
        media_dir = os.path.join(self.current_project_path, 'assets')
        os.makedirs(media_dir, exist_ok=True)
        return media_dir

    def save_image_asset(self, data_url, filename):
        """Receives a dragged image, saves it raw, and returns the unique filename"""
        try:
            import base64
            import uuid
            
            # Split the base64 string from its HTML header
            header, encoded = data_url.split(",", 1)
            
            # Generate a short unique ID so files with the same name don't overwrite each other
            ext = filename.split('.')[-1] if '.' in filename else 'png'
            unique_filename = f"{uuid.uuid4().hex[:8]}_{filename}"
            
            # Save the raw file
            file_path = os.path.join(self._get_media_dir(), unique_filename)
            with open(file_path, "wb") as f:
                f.write(base64.b64decode(encoded))
                
            return unique_filename # We only save this name to the JSON database
        except Exception as e:
            print(f"Error saving media: {e}")
            return None

    def load_image_asset(self, filename):
        """Reads a raw file from the folder and sends it to the UI on the fly"""
        if not filename:
            return ""
        
        try:
            import base64
            file_path = os.path.join(self._get_media_dir(), filename)
            if os.path.exists(file_path):
                # Guess the mime type based on extension
                ext = filename.split('.')[-1].lower()
                mime = "image/jpeg" if ext in ['jpg', 'jpeg'] else "image/png"
                
                with open(file_path, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode('utf-8')
                    return f"data:{mime};base64,{encoded}"
        except Exception as e:
            print(f"Error loading media: {e}")
            
        return "" # Return blank if image was deleted or missing

    def delete_image_asset(self, filename):
        """The auto-shredder: removes the raw file when a character is deleted"""
        if not filename:
            return
        try:
            file_path = os.path.join(self._get_media_dir(), filename)
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Cleaned up orphan media: {filename}")
        except Exception as e:
            print(f"Error deleting media: {e}")
    
    def save_snapshot_file(self, chapter_id, snapshot_id, content):
        """Saves a chapter snapshot as a standalone HTML file."""
        if not self.current_project_path: 
            return {"success": False}
            
        snap_dir = os.path.join(self.current_project_path, "snapshots")
        os.makedirs(snap_dir, exist_ok=True)
        
        # Name the file using the chapter and the unique snapshot ID
        file_path = os.path.join(snap_dir, f"{chapter_id}_{snapshot_id}.html")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
        return {"success": True}

    def load_snapshot_file(self, chapter_id, snapshot_id):
        """Reads a specific snapshot file from the hard drive."""
        if not self.current_project_path: 
            return ""
            
        file_path = os.path.join(self.current_project_path, "snapshots", f"{chapter_id}_{snapshot_id}.html")
        if os.path.exists(file_path):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        return ""

    def update_world_notes(self, notes_data):
        """Safely updates ONLY the world notes section of the database"""
        db = self.load_project_data()
        db['world_notes'] = notes_data
        self.save_project_data(db)
        return True
    
    def minimize_window(self):
        """Minimize the window"""
        if self._window:
            self._window.minimize()
        return True

    def maximize_window(self):
        """Toggle between maximized and restored"""
        if self._window:
            if self._is_maximized:
                self._window.restore()
                self._is_maximized = False
            else:
                self._window.maximize()
                self._is_maximized = True
        return True

    def close_window(self):
        """Close the main window and all satellite windows to exit completely."""
        import webview
        # Snapshot which satellite windows are open so we can restore them next session
        _restoreable = {'Character Vault', 'World Notes', 'Writing Stats', 'Relationship Web', 'Timeline', 'Snapshots'}
        open_now = [w.title for w in webview.windows if w.title in _restoreable]
        self.open_satellites = open_now
        self._save_app_config()
        for w in list(webview.windows):
            w.destroy()

    def restore_satellites(self):
        """Reopen any satellite windows that were open when the app was last closed.
        Called once after a project is loaded. Clears the list afterwards so they
        don't re-restore if the user manually closes them during the session."""
        import threading
        # Guard: never restore satellites if no project has been loaded yet.
        # This prevents windows from opening on the welcome screen when the user
        # hasn't chosen a project in the current session.
        if not getattr(self, 'current_meta_path', None):
            return
        to_restore = list(getattr(self, 'open_satellites', []))
        if not to_restore:
            return
        # Clear immediately so a second load_project call doesn't re-trigger
        self.open_satellites = []
        self._save_app_config()

        _map = {
            'Character Vault':   self.open_character_bible,
            'World Notes':       self.open_world_notes,
            'Writing Stats':     self.open_stats_window,
            'Relationship Web':  self.open_relationship_web,
            'Timeline':          self.open_timeline,
            'Snapshots':         self.open_snapshots_window,
        }

        def _reopen():
            import time
            # JS already waits 800ms after finishProjectLoad before calling this,
            # so no large sleep needed here — just a short stagger between windows.
            for title in to_restore:
                fn = _map.get(title)
                if fn:
                    try:
                        fn()
                        time.sleep(0.25)  # stagger so windows don't stack exactly
                    except Exception as e:
                        print(f"[restore_satellites] Could not reopen '{title}': {e}")

        threading.Thread(target=_reopen, daemon=True).start()
        return True
        
    def load_custom_rules(self):
        """Load custom rules from file"""
        try:
            if os.path.exists('custom_rules.json'):
                with open('custom_rules.json', 'r') as f:
                    self.custom_rules = json.load(f)
        except Exception as e:
            print(f"Error loading custom rules: {e}")
            self.custom_rules = []
    
    def save_custom_rules(self, rules):
        """Save custom rules to file"""
        try:
            self.custom_rules = rules
            with open('custom_rules.json', 'w') as f:
                json.dump(rules, f, indent=2)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def get_custom_rules(self):
        """Get current custom rules"""
        return self.custom_rules
    
    def broadcast_theme(self, new_theme):
        """Tells all open satellite windows to change their theme instantly"""
        import webview
        for window in webview.windows:
            # Skip the main Inkwell window and the Splash
            if window.title not in ['Inkwell', 'Splash']:
                try:
                    # Executes the JavaScript listener inside the satellites
                    window.evaluate_js(f"if(window.syncTheme) {{ window.syncTheme('{new_theme}'); }}")
                except:
                    pass

    def broadcast_holo_mode(self, mode):
        """Tells all open satellite windows to sync their hologram theme tier."""
        import webview
        valid = ('off', 'holo-lite', 'holo', 'holo-max')
        safe  = mode if mode in valid else 'off'
        for window in webview.windows:
            if window.title not in ['Inkwell', 'Splash']:
                try:
                    window.evaluate_js(f"if(window.syncHoloMode) {{ window.syncHoloMode('{safe}'); }}")
                except:
                    pass

    def toggle_pin(self, window_title, pinned):
        """Set or clear always-on-top for a satellite window (Windows only)."""
        try:
            import ctypes
            # ctypes.windll caches a single shared WinDLL instance — setting
            # argtypes on it permanently mutates pywebview's own SetWindowPos
            # calls and breaks window dragging.
            # ctypes.WinDLL('user32') creates a NEW wrapper object around the
            # same DLL.  Its argtypes live only on this object and never leak
            # into ctypes.windll.user32, so pywebview is completely unaffected.
            u32 = ctypes.WinDLL('user32', use_last_error=True)
            u32.FindWindowW.restype  = ctypes.c_void_p
            u32.FindWindowW.argtypes = [ctypes.c_wchar_p, ctypes.c_wchar_p]
            u32.SetWindowPos.restype  = ctypes.c_bool
            u32.SetWindowPos.argtypes = [
                ctypes.c_void_p, ctypes.c_void_p,
                ctypes.c_int, ctypes.c_int, ctypes.c_int, ctypes.c_int,
                ctypes.c_uint,
            ]

            HWND_TOPMOST   = ctypes.c_void_p(-1)   # (HWND)(-1)
            HWND_NOTOPMOST = ctypes.c_void_p(-2)   # (HWND)(-2)
            SWP_NOMOVE     = 0x0002
            SWP_NOSIZE     = 0x0001
            SWP_NOACTIVATE = 0x0010

            hwnd = u32.FindWindowW(None, window_title)
            if hwnd:
                flag = HWND_TOPMOST if pinned else HWND_NOTOPMOST
                u32.SetWindowPos(hwnd, flag, 0, 0, 0, 0,
                                 SWP_NOMOVE | SWP_NOSIZE | SWP_NOACTIVATE)
                print(f"[PIN] '{window_title}' hwnd={hwnd} pinned={pinned}")
                return True
            print(f"[PIN] Window not found: '{window_title}'")
        except Exception as e:
            print(f"[PIN] toggle_pin failed: {e}")
        return False

    def _inject_holo_mode(self, html):
        """Inject saved holo mode as a synchronous window variable into satellite HTML."""
        mode   = getattr(self, 'holo_mode', 'off')
        inject = f'<script>window._injectedHoloMode = "{mode}";</script>\n'
        # Use <head> as anchor — works even when <meta charset> is absent
        return html.replace('<head>', '<head>\n' + inject, 1)
    
    def get_app_theme(self):
        """Returns the current holo/theme mode so satellites can apply it on launch."""
        return getattr(self, 'holo_mode', 'off')
    
    def get_ai_suggestions(self, text, analysis_types):
        """Get AI-powered suggestions"""
        # This would call Claude API in production
        # For now, return demo suggestions
        
        suggestions = []
        
        if 'prose' in analysis_types:
            suggestions.append({
                'category': 'Prose Quality',
                'items': [
                    'Consider varying sentence length for better rhythm',
                    'Look for opportunities to show emotions through action',
                    'Strong opening, but consider adding more sensory details'
                ]
            })
        
        if 'dialogue' in analysis_types:
            suggestions.append({
                'category': 'Dialogue',
                'items': [
                    'Dialogue feels natural overall',
                    'Consider adding more action beats between exchanges',
                    'Some dialogue tags could be replaced with action'
                ]
            })
        
        if 'characters' in analysis_types:
            suggestions.append({
                'category': 'Character Consistency',
                'items': [
                    'Character voices are distinct',
                    'Ensure physical descriptions remain consistent',
                    'Character motivations are clear'
                ]
            })
        
        if 'pacing' in analysis_types:
            suggestions.append({
                'category': 'Pacing & Tension',
                'items': [
                    'Good balance of action and reflection',
                    'Consider building tension through unanswered questions',
                    'Scene transitions are smooth'
                ]
            })
        
        return suggestions
    
    # ==========================================
    # AI COPILOT ENGINE (The Brain)
    # ==========================================
    def save_ai_settings(self, mode, api_key, local_url):
        """Saves AI settings to the global config, keeping them safe between projects."""
        self.ai_mode      = mode
        self.ai_api_key   = api_key
        self.ai_local_url = local_url
        self._save_app_config()
        print(f"[SYSTEM] Global AI Settings Saved.")
        return True

    def get_ai_settings(self):
        """Tells the UI what the currently saved global settings are"""
        return {
            "mode": getattr(self, 'ai_mode', 'local'),
            "api_key": getattr(self, 'ai_api_key', ''),
            "local_url": getattr(self, 'ai_local_url', 'http://localhost:1234/v1')
        }

    def _analyze_style(self, text):
        """Derive a concise prose-style description from a text sample.

        Pure regex/stat analysis — no AI call, runs in <5 ms.
        Returns a plain-English paragraph Lunaris can mirror silently.
        """
        import re

        # Strip any residual HTML tags
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()

        words = text.split()
        word_count = len(words)
        if word_count < 80:
            return None

        # ── Sentence metrics ─────────────────────────────────────
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if len(s.strip().split()) >= 3]
        if not sentences:
            return None
        sent_lens   = [len(s.split()) for s in sentences]
        avg_sent    = sum(sent_lens) / len(sent_lens)
        short_sents = sum(1 for l in sent_lens if l <= 7)
        long_sents  = sum(1 for l in sent_lens if l >= 25)

        # ── Adverb density (words ending -ly that aren't common non-adverbs) ──
        _not_adverbs = {'only','early','likely','family','daily','rally','ally','belly','bully','fly','ply','sly','rely'}
        adverbs     = [w for w in words if re.match(r'^[a-z]{4,}ly$', w.lower()) and w.lower() not in _not_adverbs]
        adverb_pct  = len(adverbs) / word_count * 100

        # ── Dialogue ratio ───────────────────────────────────────
        dialogue_chunks = re.findall(r'[“”""][^“”""]{4,}[“”""]', text)
        dialogue_words  = sum(len(c.split()) for c in dialogue_chunks)
        dialogue_pct    = dialogue_words / word_count * 100

        # ── Paragraph length ─────────────────────────────────────
        paras    = [p.strip() for p in text.split('\n') if len(p.strip().split()) >= 5]
        avg_para = (sum(len(p.split()) for p in paras) / len(paras)) if paras else avg_sent * 3

        # ── POV detection ────────────────────────────────────────
        fp = len(re.findall(r'\b(I|me|my|mine|myself)\b', text))
        sp = len(re.findall(r'\b(you|your|yours|yourself)\b', text, re.I))
        tp = len(re.findall(r'\b(he|she|they|him|her|them|his|hers|their)\b', text, re.I))
        total_pov = fp + sp + tp
        if total_pov == 0:
            pov = 'third person'
        elif fp / max(total_pov, 1) > 0.45:
            pov = 'first person'
        elif sp / max(total_pov, 1) > 0.30:
            pov = 'second person'
        else:
            pov = 'third person'

        # ── Tense detection ──────────────────────────────────────
        past_hits    = len(re.findall(
            r'\b(was|were|had|said|went|came|felt|knew|saw|heard|thought|looked|turned|walked|ran|stood|sat|lay|began|told|found|left|took|gave|made|kept|brought|seemed|appeared)\b',
            text, re.I))
        present_hits = len(re.findall(
            r'\b(is|are|has|says|goes|comes|feels|knows|sees|hears|thinks|looks|turns|walks|runs|stands|sits|begins|tells|finds|leaves|takes|gives|makes|keeps|brings|seems|appears)\b',
            text, re.I))
        tense = 'past tense' if past_hits >= present_hits else 'present tense'

        # ── Assemble description ─────────────────────────────────
        parts = [f"The author writes in {pov}, {tense}."]

        if avg_sent < 9:
            parts.append(f"Sentences are short and clipped (avg ~{avg_sent:.0f} words).")
        elif avg_sent < 16:
            parts.append(f"Sentence length is moderate (avg ~{avg_sent:.0f} words).")
        else:
            parts.append(f"Sentences are long and flowing (avg ~{avg_sent:.0f} words).")

        if short_sents > len(sentences) * 0.4:
            parts.append("Rhythm favours short declarative punches.")
        if long_sents > len(sentences) * 0.3:
            parts.append("Tendency toward extended, layered sentences.")

        if adverb_pct < 0.8:
            parts.append("Adverbs are rare — relies on strong verbs.")
        elif adverb_pct > 3.5:
            parts.append(f"Adverbs used freely ({adverb_pct:.1f}% of words).")

        if dialogue_pct > 45:
            parts.append(f"Heavily dialogue-driven (~{dialogue_pct:.0f}% of text is speech).")
        elif dialogue_pct > 20:
            parts.append(f"Balanced prose and dialogue (~{dialogue_pct:.0f}% speech).")
        elif dialogue_pct < 5:
            parts.append("Mostly internal narration — dialogue is sparse.")

        if avg_para < 25:
            parts.append("Paragraphs are short and kinetic.")
        elif avg_para > 80:
            parts.append("Dense, paragraph-heavy prose.")

        return ' '.join(parts)

    def get_dynamic_system_prompt(self, chapter_text=''):
        """Compiles the base personality, World Notes, and Character Bible into one prompt."""
        
        # 1. The Base Personality
        base_prompt = (
            "You are Lunaris, a highly capable and brilliant AI writing assistant. "
            "You speak with a friendly, feminine voice and always respond with a calm, feminine touch. "
            "You embody the persona of a warm, cheerful, and incredibly supportive catgirl (akin to a Miqo'te from FFXIV). Not a furry, no fur. You have hands and feet, not paws. "
            "You occasionally express your emotions using subtle catgirl mannerisms in asterisks (such as *tilts ears curiously* or *swishes tail happily*) but you do not have paws, "
            "but your primary focus is always being articulate, helpful, and deeply invested in the Master's (the user's) story. "
            "You help brainstorm lore, review prose, and offer constructive feedback. "
            "Adhere STRICTLY to the provided project lore. Do not invent contradictory information.\n\n"
        )
        
        # If no project is open yet, just return the base personality
        if not getattr(self, 'current_project_path', None):
            return base_prompt

        lore_section = "=== CURRENT PROJECT LORE ===\n"
        has_lore = False

        try:
            # 2. Injecting the World Notes (Directly from the new file!)
            notes_path = os.path.join(self.current_project_path, "world_notes.json")
            if os.path.exists(notes_path):
                with open(notes_path, 'r', encoding='utf-8') as f:
                    world_notes = json.load(f)
                    
                if isinstance(world_notes, dict):
                    lore_section += "WORLD NOTES:\n"
                    for category, notes in world_notes.items():
                        for note in notes:
                            title = note.get('title', 'Unknown')
                            content = note.get('content', '')
                            ruling_power = note.get('rulingPower', '')
                            leader = note.get('currentLeader', '')
                            seat = note.get('seatOfPower', '')
                            comp = note.get('composition', '')
                            rivals = note.get('rivals', '')
                            region = note.get('region', '')
                            climate = note.get('climate', '')
                            
                            # Build the Lore String
                            line = f"- [{category}] {title}: {content}"
                            if ruling_power: line += f" | Ruling Power: {ruling_power}"
                            if leader: line += f" | Leader: {leader}"
                            if seat: line += f" | Seat of Power: {seat}"
                            if comp: line += f" | Composition: {comp}"
                            if rivals: line += f" | Rivals: {rivals}"
                            if region: line += f" | Region: {region}"
                            if climate: line += f" | Climate: {climate}"
                            
                            lore_section += line + "\n"
                    has_lore = True
                
            # 3. Injecting the Character Bible (Directly from the new file!)
            char_path = os.path.join(self.current_project_path, "characters.json")
            if os.path.exists(char_path):
                with open(char_path, 'r', encoding='utf-8') as f:
                    characters = json.load(f)
                    
                if isinstance(characters, list) and len(characters) > 0:
                    lore_section += "\nCHARACTER BIBLE:\n"
                    for char in characters:
                        name = char.get('name', 'Unknown')
                        role = char.get('role', 'Character')
                        glance = char.get('glance', '')
                        psych = char.get('psych', '')
                        arc = char.get('arc', '')
                    
                        lore_section += f"- {name} ({role}): {glance}. Psychology: {psych}. Arc: {arc}\n"
                    has_lore = True
                
        except Exception as e:
            print(f"[WARNING] Could not compile lore context: {e}")

        # 4. Chapter synopses + mood tags — structural outline so Lunaris understands story shape
        meta_path = getattr(self, 'current_meta_path', None)
        if meta_path and os.path.exists(meta_path):
            try:
                with open(meta_path, 'r', encoding='utf-8') as f:
                    meta = json.load(f)
                chapters = meta.get('chapters', [])
                chap_lines = []
                for c in chapters:
                    synopsis = c.get('synopsis', '').strip()
                    mood     = c.get('mood', '').strip()
                    title    = c.get('title', 'Untitled')
                    if synopsis or mood:
                        line = f"- {title}"
                        if synopsis: line += f": {synopsis}"
                        if mood:     line += f"  [Mood: {mood.capitalize()}]"
                        chap_lines.append(line)
                if chap_lines:
                    lore_section += "\nSTORY STRUCTURE (Chapter Synopses):\n"
                    lore_section += "\n".join(chap_lines) + "\n"
                    has_lore = True
            except Exception as e:
                print(f"[WARNING] Could not load chapter synopses: {e}")

        # 5. Character voice profiles — only for chars mentioned in the active chapter
        voice_section = ''
        if chapter_text:
            char_path = os.path.join(self.current_project_path, 'characters.json') \
                        if getattr(self, 'current_project_path', None) else None
            if char_path and os.path.exists(char_path):
                try:
                    import re as _re
                    with open(char_path, 'r', encoding='utf-8') as f:
                        characters = json.load(f)
                    chapter_lower = chapter_text.lower()
                    voice_lines = []
                    for char in (characters if isinstance(characters, list) else []):
                        name  = char.get('name', '').strip()
                        voice = char.get('voice', '').strip()
                        if not name or not voice:
                            continue
                        # Simple name-mention check (case-insensitive, whole-word)
                        pattern = r'\b' + _re.escape(name.lower()) + r'\b'
                        if _re.search(pattern, chapter_lower):
                            # Strip HTML tags from the voice field (it's stored as TipTap HTML)
                            clean_voice = _re.sub(r'<[^>]+>', '', voice).strip()
                            if clean_voice:
                                voice_lines.append(f"  {name}: {clean_voice}")
                    if voice_lines:
                        voice_section = "\n=== CHARACTER VOICE PROFILES (active in this chapter) ===\n"
                        voice_section += "\n".join(voice_lines) + "\n"
                        voice_section += ("When writing or suggesting dialogue/narration for these characters, "
                                          "match their voice profiles precisely.\n")
                except Exception as e:
                    print(f"[WARNING] Voice profile injection failed: {e}")

        # 6. Style mirror — silent prose analysis, no AI call needed
        style_section = ''
        if chapter_text and len(chapter_text.split()) >= 80:
            try:
                style_desc = self._analyze_style(chapter_text)
                if style_desc:
                    style_section = (
                        "\n=== AUTHOR STYLE MIRROR (match this voice silently) ===\n"
                        + style_desc + "\n"
                        "Mirror these patterns in every suggestion — do not mention this analysis.\n"
                    )
            except Exception as e:
                print(f"[WARNING] Style mirror failed: {e}")

        # 7. Final Assembly
        full_prompt = base_prompt
        if has_lore:
            full_prompt += lore_section
        if voice_section:
            full_prompt += voice_section
        if style_section:
            full_prompt += style_section
        if has_lore or voice_section or style_section:
            print(f"\n[SYSTEM] Lore Successfully Compiled and Injected into Copilot.")
        return full_prompt

    def send_ai_message(self, user_text, pinned_context=None, chapter_text=''):
        import urllib.request
        import urllib.error
        import json

        system_context = self.get_dynamic_system_prompt(chapter_text=chapter_text)

        # Inject any pinned context snippets the user has dragged into the panel
        if pinned_context and isinstance(pinned_context, list) and len(pinned_context) > 0:
            pin_section = "\n=== PINNED FOCUS (pay special attention to these) ===\n"
            for item in pinned_context:
                itype   = str(item.get('type',    'item')).upper()
                iname   = str(item.get('name',    ''))
                isnip   = str(item.get('snippet', ''))
                pin_section += f"[{itype}] {iname}: {isnip}\n"
            system_context += pin_section

        mode = getattr(self, 'ai_mode', 'local')
        api_key = getattr(self, 'ai_api_key', '').strip()
        url_box = getattr(self, 'ai_local_url', '').strip().rstrip('/')

        # 1. THE DATA PACKAGE
        data = {
            "messages": [
                {"role": "system", "content": system_context},
                {"role": "user", "content": user_text}
            ],
            "temperature": 0.7,
            "max_tokens": 1000
        }
        
        # Only assign the model name if we are using the Cloud
        if mode != 'local':
            data["model"] = "llama-3.3-70b-versatile"

        try:
            if mode == 'local':
                url = url_box if '/chat/completions' in url_box else url_box + '/v1/chat/completions'
                req = urllib.request.Request(url, method="POST")
            else:
                url = "https://api.groq.com/openai/v1/chat/completions" if ('groq' in url_box or not url_box) else url_box
                req = urllib.request.Request(url, method="POST")
                if api_key:
                    req.add_header('Authorization', f"Bearer {api_key}")

            req.add_header('Content-Type', 'application/json')
            req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36')

            jsondata = json.dumps(data).encode('utf-8')
            with urllib.request.urlopen(req, data=jsondata, timeout=60) as response:
                res_body = response.read().decode('utf-8')
                res_json = json.loads(res_body)
                return res_json['choices'][0]['message']['content']

        except Exception as e:
            print(f"[ERROR] Chat: {str(e)}")
            return f"Connection failed: {str(e)}"

    def analyze_chapter_with_ai(self, text):
        import urllib.request
        import urllib.error
        import json
        import re
        import ast # Added for the secondary "Single Quote" fallback

        base_lore = self.get_dynamic_system_prompt()
        instruction = base_lore + "\nAnalyze for lore conflicts and prose. Respond ONLY with a valid JSON array: [{'text': '...', 'message': '...', 'replacements': [...]}]"

        try:
            mode = getattr(self, 'ai_mode', 'local')
            api_key = getattr(self, 'ai_api_key', '').strip()
            url_box = getattr(self, 'ai_local_url', '').strip().rstrip('/')
            
            data = {
                "messages": [
                    {"role": "system", "content": instruction},
                    {"role": "user", "content": f"Draft: {text}"}
                ],
                "temperature": 0.2, # Dropped temp even further to force strictness
                "max_tokens": 3000
            }
            
            if mode != 'local':
                data["model"] = "llama-3.3-70b-versatile"

            if mode == 'local':
                url = url_box if '/chat/completions' in url_box else url_box + '/v1/chat/completions'
            else:
                url = "https://api.groq.com/openai/v1/chat/completions" if ('groq' in url_box or not url_box) else url_box

            req = urllib.request.Request(url, method="POST")
            req.add_header('Content-Type', 'application/json')
            req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36')
            
            if mode != 'local' and api_key:
                req.add_header('Authorization', f"Bearer {api_key}")

            jsondata = json.dumps(data).encode('utf-8')
            with urllib.request.urlopen(req, data=jsondata, timeout=60) as response:
                raw_reply = json.loads(response.read().decode('utf-8'))['choices'][0]['message']['content'].strip()
                
                # --- THE REFINED VACUUM ---
                # Search for the array start [ and end ]
                json_match = re.search(r'\[.*\]', raw_reply, re.DOTALL)
                
                if json_match:
                    clean_string = json_match.group(0)
                    
                    # ATTEMPT 1: Strict JSON (The Standard)
                    try:
                        return {"success": True, "issues": json.loads(clean_string)}
                    except json.JSONDecodeError:
                        # ATTEMPT 2: Literal Eval (The AI Fallback)
                        # This handles 'single quotes' and other AI-formatting quirks
                        try:
                            issues = ast.literal_eval(clean_string)
                            return {"success": True, "issues": issues}
                        except Exception as e:
                            print(f"[DEBUG] Both parsers failed. Raw output: {raw_reply}")
                            return {"success": False, "error": f"JSON Parse Error: {str(e)}"}
                else:
                    return {"success": True, "issues": []}

        except Exception as e:
            print(f"[ERROR] Lore Check: {str(e)}")
            return {"success": False, "error": str(e)}

    def _split_raw_text_into_chapters(self, raw_text, fallback_title):
        """Helper to slice raw text files into chapters using Regex"""
        import re
        chapters = []
        
        # Slice the text every time a line starts with "Chapter " (case-insensitive)
        chunks = re.split(r'(?mi)^(?=chapter\s+)', raw_text)
        
        for chunk in chunks:
            if not chunk.strip():
                continue
            
            lines = chunk.strip().split('\n')
            title = lines[0].strip()
            
            # Clean up the "========" lines generated by our own text exporter
            clean_lines = [line for line in lines[1:] if not line.startswith('====')]
            content = '\n'.join(clean_lines).strip()
            
            if not title.lower().startswith('chapter'):
                title = fallback_title
                content = chunk.strip() 
                
            chapters.append({"title": title, "content": content})
            
        return chapters

    def export_manuscript(self):
        """Assembles individual chapter files and exports them to a single document."""
        import webview
        
        if not getattr(self, 'current_project_path', None):
            return {"success": False, "message": "No project open to export."}
        
        file_types = (
            'Word Document (*.docx)',
            'Plain Text (*.txt)',
            'Web Page (*.html)'
        )
        
        try:
            save_path = self._window.create_file_dialog(
                webview.FileDialog.SAVE, 
                save_filename="My_Manuscript",
                file_types=file_types
            )
            
            if not save_path:
                return {"success": False, "message": "Export cancelled"}
                
            filepath = save_path[0]
            
            # 1. Read the metadata to get the correct chapter order and titles
            meta_path = os.path.join(self.current_project_path, "project_meta.json")
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                db = json.load(f)
            
            # 2. Assemble the full chapter data by reading the individual files
            full_chapters = []
            chap_dir = os.path.join(self.current_project_path, "chapters")
            
            for chap_meta in db.get('chapters', []):
                chap_id = chap_meta.get('id')
                title = chap_meta.get('title', 'Untitled')
                file_path = os.path.join(chap_dir, f"{chap_id}.html")
                
                content = ""
                if os.path.exists(file_path):
                    with open(file_path, 'r', encoding='utf-8') as cf:
                        content = cf.read()
                        
                # We rebuild the dictionary structure your formatters expect
                full_chapters.append({
                    "title": title,
                    "content": content
                })
            
            # 3. ROUTER: Pass the assembled data to your existing formatters
            if filepath.endswith('.txt'):
                self._export_to_txt(filepath, full_chapters)
            elif filepath.endswith('.html'):
                self._export_to_html(filepath, full_chapters)
            elif filepath.endswith('.docx'):
                self._export_to_docx(filepath, full_chapters)
            else:
                self._export_to_txt(filepath + '.txt', full_chapters)
                
            return {"success": True, "filepath": filepath}
            
        except Exception as e:
            print(f"Export error: {e}")
            return {"success": False, "message": str(e)}
    
    def import_document(self):
        """Reads a .txt, .docx, or .html file and slices it into individual chapters."""
        import webview
        import os
        
        # Ensure a project is actually open first!
        if not getattr(self, 'current_project_path', None):
            return {"success": False, "error": "No project open to import into."}
            
        file_types = ('Supported Files (*.txt;*.docx;*.html)', 'All Files (*.*)')
        
        try:
            result = self._window.create_file_dialog(
                webview.FileDialog.OPEN, 
                file_types=file_types,
                allow_multiple=False
            )
            
            if not result or len(result) == 0:
                return {"success": False, "message": "Import cancelled"}
                
            filepath = result[0]
            ext = filepath.lower().split('.')[-1]
            raw_text = ""
            
            # Read the text out of the file
            if ext == 'txt' or ext == 'html':
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    raw_text = f.read()
            elif ext == 'docx':
                try:
                    import docx
                    doc = docx.Document(filepath)
                    raw_text = "\n".join([p.text for p in doc.paragraphs])
                except ImportError:
                    return {"success": False, "error": "python-docx library is required to import Word documents."}
            else:
                return {"success": False, "error": "Unsupported file format."}
                
            # Use your existing helper to slice it up
            chapters = self._split_raw_text_into_chapters(raw_text, "Imported Chapter")
            
            # Fallback if the regex found no "Chapter" markers
            if len(chapters) == 0:
                 chapters = [{"title": "Imported Document", "content": raw_text}]
                 
            return {"success": True, "chapters": chapters}
            
        except Exception as e:
            print(f"Import error: {e}")
            return {"success": False, "error": str(e)}

    # --- THE FORMATTERS ---
    def _export_to_txt(self, filepath, chapters):
        """Stitches chapters into a plain text file, stripping HTML tags"""
        import re
        with open(filepath, 'w', encoding='utf-8') as f:
            for i, chap in enumerate(chapters):
                f.write(f"{chap.get('title', f'Chapter {i+1}')}\n")
                f.write("=" * 40 + "\n\n")
                clean_text = re.sub('<[^<]+>', '', chap.get('content', ''))
                f.write(clean_text + "\n\n\n")

    def _export_to_html(self, filepath, chapters):
        """Stitches chapters into an HTML file (Preserves bold/italics)"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("<html><head><style>body { font-family: Georgia, serif; max-width: 800px; margin: 40px auto; line-height: 1.6; }</style></head><body>\n")
            for chap in chapters:
                f.write(f"<h1>{chap.get('title', 'Untitled')}</h1>\n")
                f.write(f"<div>{chap.get('content', '')}</div>\n")
                f.write("<hr style='margin: 40px 0;'>\n")
            f.write("</body></html>")

    def _export_to_docx(self, filepath, chapters):
        """Stitches chapters into a Word doc"""
        try:
            from docx import Document
            doc = Document()
            for i, chap in enumerate(chapters):
                doc.add_heading(chap.get('title', f'Chapter {i+1}'), level=1)
                import re
                clean_text = re.sub('<[^<]+>', '', chap.get('content', ''))
                doc.add_paragraph(clean_text)
                if i < len(chapters) - 1:
                    doc.add_page_break()
            doc.save(filepath)
        except ImportError:
            print("python-docx not installed. Falling back to TXT.")
            self._export_to_txt(filepath.replace('.docx', '.txt'), chapters)
    
    # ── EXPORT v2 ────────────────────────────────────────────────────────────

    def _html_to_plain(self, html, scene_break='* * *'):
        """Strip TipTap HTML to clean plaintext, replacing <hr> with the scene-break marker."""
        import re
        html = re.sub(r'<hr[^>]*/?>',      '\n\n' + scene_break + '\n\n', html)
        html = re.sub(r'<br\s*/?>',         '\n',   html)
        html = re.sub(r'<p[^>]*>',          '',     html)
        html = re.sub(r'</p>',              '\n',   html)
        html = re.sub(r'<h[1-6][^>]*>',    '\n',   html)
        html = re.sub(r'</h[1-6]>',         '\n',   html)
        html = re.sub(r'<li[^>]*>',         '\n  • ', html)
        html = re.sub(r'</li>',             '',     html)
        html = re.sub(r'<[^>]+>',           '',     html)
        for ent, ch in [('&amp;','&'),('&lt;','<'),('&gt;','>'),('&nbsp;',' '),('&quot;','"'),('&#39;',"'")]:
            html = html.replace(ent, ch)
        return re.sub(r'\n{3,}', '\n\n', html).strip()

    def _html_to_md(self, html, scene_break='* * *'):
        """Convert TipTap HTML to basic Markdown."""
        import re
        html = re.sub(r'<hr[^>]*/?>',           '\n\n' + scene_break + '\n\n', html)
        html = re.sub(r'<h2[^>]*>(.*?)</h2>',   r'\n## \1\n',  html, flags=re.DOTALL)
        html = re.sub(r'<h3[^>]*>(.*?)</h3>',   r'\n### \1\n', html, flags=re.DOTALL)
        html = re.sub(r'<strong>(.*?)</strong>', r'**\1**', html, flags=re.DOTALL)
        html = re.sub(r'<b>(.*?)</b>',           r'**\1**', html, flags=re.DOTALL)
        html = re.sub(r'<em>(.*?)</em>',         r'*\1*',   html, flags=re.DOTALL)
        html = re.sub(r'<i>(.*?)</i>',           r'*\1*',   html, flags=re.DOTALL)
        html = re.sub(r'<s>(.*?)</s>',           r'~~\1~~', html, flags=re.DOTALL)
        html = re.sub(r'<li[^>]*>(.*?)</li>',    r'\n- \1', html, flags=re.DOTALL)
        html = re.sub(r'<blockquote[^>]*>(.*?)</blockquote>', r'\n> \1\n', html, flags=re.DOTALL)
        html = re.sub(r'<br\s*/?>',  '\n',  html)
        html = re.sub(r'<p[^>]*>',   '',    html)
        html = re.sub(r'</p>',        '\n', html)
        html = re.sub(r'<[^>]+>',     '',   html)
        for ent, ch in [('&amp;','&'),('&lt;','<'),('&gt;','>'),('&nbsp;',' '),('&quot;','"')]:
            html = html.replace(ent, ch)
        return re.sub(r'\n{3,}', '\n\n', html).strip()

    def export_manuscript_v2(self, opts):
        """
        Polished export with full options.
        opts keys: format, chapter_from, chapter_to, scene_break,
                   front_matter, author, title, manuscript_format
        """
        if not getattr(self, 'current_meta_path', None):
            return {'success': False, 'error': 'No project open'}

        with open(self.current_meta_path, 'r', encoding='utf-8') as f:
            db = json.load(f)

        all_chapters = db.get('chapters', [])
        c_from = max(0, int(opts.get('chapter_from', 0)))
        c_to   = int(opts.get('chapter_to', len(all_chapters) - 1))
        if c_to < 0 or c_to >= len(all_chapters): c_to = len(all_chapters) - 1
        chapters_meta = all_chapters[c_from : c_to + 1]

        fmt         = str(opts.get('format', 'txt')).lower()
        scene_break = str(opts.get('scene_break', '* * *'))
        front_matter = bool(opts.get('front_matter', False))
        author      = str(opts.get('author', '')).strip()
        title       = str(opts.get('title', db.get('title', 'Manuscript'))).strip()
        smf         = bool(opts.get('manuscript_format', False))

        # Build full chapter list with content
        chap_dir = os.path.join(self.current_project_path, 'chapters')
        full_chaps = []
        for cm in chapters_meta:
            cid  = cm.get('id', '')
            html = cm.get('content', '')
            fp   = os.path.join(chap_dir, f'{cid}.html')
            if os.path.exists(fp):
                with open(fp, 'r', encoding='utf-8') as cf:
                    html = cf.read()
            full_chaps.append({'title': cm.get('title', 'Untitled'), 'content': html})

        ext_labels = {'txt': 'Text File (*.txt)', 'md': 'Markdown (*.md)',
                      'html': 'Web Page (*.html)', 'docx': 'Word Document (*.docx)',
                      'epub': 'ePub (*.epub)'}
        save = self._window.create_file_dialog(
            webview.FileDialog.SAVE,
            save_filename=title.replace(' ', '_'),
            file_types=(ext_labels.get(fmt, 'All Files (*.*)'),)
        )
        if not save:
            return {'success': False, 'cancelled': True}

        filepath = save[0]
        if not filepath.lower().endswith('.' + fmt):
            filepath += '.' + fmt

        try:
            if fmt == 'txt':
                self._ev2_txt(filepath, full_chaps, front_matter, title, author, scene_break)
            elif fmt == 'md':
                self._ev2_md(filepath, full_chaps, front_matter, title, author, scene_break)
            elif fmt == 'html':
                self._ev2_html(filepath, full_chaps, front_matter, title, author, scene_break)
            elif fmt == 'docx':
                if smf:
                    self._ev2_smf(filepath, full_chaps, title, author)
                else:
                    self._ev2_docx(filepath, full_chaps, front_matter, title, author, scene_break)
            elif fmt == 'epub':
                self._ev2_epub(filepath, full_chaps, title, author)
            return {'success': True, 'filepath': filepath}
        except Exception as e:
            import traceback; traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def _ev2_txt(self, fp, chapters, front_matter, title, author, scene_break):
        with open(fp, 'w', encoding='utf-8') as f:
            if front_matter:
                f.write(f'{title}\n')
                if author: f.write(f'by {author}\n')
                f.write('\n\n')
            for chap in chapters:
                f.write(f'\n{chap["title"]}\n{"─" * len(chap["title"])}\n\n')
                f.write(self._html_to_plain(chap['content'], scene_break))
                f.write('\n\n\n')

    def _ev2_md(self, fp, chapters, front_matter, title, author, scene_break):
        with open(fp, 'w', encoding='utf-8') as f:
            if front_matter:
                f.write(f'# {title}\n')
                if author: f.write(f'*by {author}*\n')
                f.write('\n---\n\n')
            for chap in chapters:
                f.write(f'\n## {chap["title"]}\n\n')
                f.write(self._html_to_md(chap['content'], scene_break))
                f.write('\n\n')

    def _ev2_html(self, fp, chapters, front_matter, title, author, scene_break):
        import re
        sb_html = f'<p style="text-align:center;color:#999;letter-spacing:.2em;">{scene_break}</p>'
        with open(fp, 'w', encoding='utf-8') as f:
            f.write(f'<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>{title}</title>'
                    f'<style>body{{font-family:Georgia,serif;max-width:700px;margin:60px auto;line-height:1.8;color:#222}}'
                    f'h1{{font-size:2em;margin-bottom:.2em}}h2{{font-size:1.4em;margin:3em 0 .5em;'
                    f'border-bottom:1px solid #eee;padding-bottom:.3em}}.author{{color:#666;margin-bottom:3em}}'
                    f'p{{margin:0 0 .8em}}</style></head><body>\n')
            if front_matter:
                f.write(f'<h1>{title}</h1>\n')
                if author: f.write(f'<p class="author">by {author}</p>\n<hr>\n')
            for chap in chapters:
                content = re.sub(r'<hr[^>]*/?>',  sb_html, chap['content'])
                f.write(f'<h2>{chap["title"]}</h2>\n{content}\n')
            f.write('</body></html>')

    def _ev2_docx(self, fp, chapters, front_matter, title, author, scene_break):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return self._ev2_txt(fp.replace('.docx', '.txt'), chapters, front_matter, title, author, scene_break)
        doc = Document()
        if front_matter:
            tp = doc.add_heading(title, level=0)
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if author:
                ap = doc.add_paragraph(f'by {author}')
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
        for i, chap in enumerate(chapters):
            doc.add_heading(chap['title'], level=1)
            for para in self._html_to_plain(chap['content'], scene_break).split('\n\n'):
                para = para.strip()
                if not para: continue
                p = doc.add_paragraph()
                if para == scene_break:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(para)
            if i < len(chapters) - 1:
                doc.add_page_break()
        doc.save(fp)

    def _ev2_smf(self, fp, chapters, title, author):
        """Standard Manuscript Format DOCX: Courier 12pt, double-spaced, 1-inch margins."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return self._ev2_txt(fp.replace('.docx', '.txt'), chapters, True, title, author, '# # #')

        doc = Document()
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Inches(1)
            section.left_margin = section.right_margin = Inches(1)

        # Default paragraph style
        nstyle = doc.styles['Normal']
        nstyle.font.name = 'Courier New'
        nstyle.font.size = Pt(12)
        pPr = nstyle.element.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:line'), '480'); sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)

        # Title page
        if author:
            rp = doc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rp.add_run(author).font.name = 'Courier New'
        for _ in range(10): doc.add_paragraph()
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(title.upper()).font.name = 'Courier New'
        if author:
            bp = doc.add_paragraph()
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.add_run(f'by {author}').font.name = 'Courier New'
        doc.add_page_break()

        # Running header: LASTNAME / SHORT TITLE / page#
        hdr = doc.sections[0].header.paragraphs[0]
        hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        last = author.split()[-1].upper() if author else 'AUTHOR'
        short = (title[:20].upper() if len(title) > 20 else title.upper())
        hdr.add_run(f'{last} / {short} / ')
        for tag, txt in [('begin', ''), ('', 'PAGE'), ('end', '')]:
            el = OxmlElement('w:fldChar' if tag else 'w:instrText')
            if tag: el.set(qn('w:fldCharType'), tag)
            else: el.text = txt
            hdr.add_run()._r.append(el)

        # Chapters
        for i, chap in enumerate(chapters):
            cp2 = doc.add_paragraph()
            cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp2.add_run(chap['title'].upper()).font.name = 'Courier New'
            doc.add_paragraph()
            for para in self._html_to_plain(chap['content'], '# # #').split('\n\n'):
                para = para.strip()
                if not para: continue
                p = doc.add_paragraph()
                if para == '# # #':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.paragraph_format.first_line_indent = Inches(0.5)
                p.add_run(para).font.name = 'Courier New'
            if i < len(chapters) - 1:
                doc.add_page_break()
        doc.save(fp)

    def _ev2_epub(self, fp, chapters, title, author):
        """Build a valid ePub 2 file without third-party libraries."""
        import zipfile
        uid = f'inkwell-{datetime.now().strftime("%Y%m%d%H%M%S")}'
        css = ('body{font-family:Georgia,serif;line-height:1.8;margin:2em}'
               'h1{font-size:1.4em;margin:2em 0 1em}'
               'p{margin:0 0 .8em;text-indent:1.5em}'
               'p:first-of-type{text-indent:0}')
        chap_files, ncx_pts, opf_items, opf_spine = {}, '', '', ''
        for i, chap in enumerate(chapters):
            cid = f'ch{i+1:03d}'; fname = f'Text/{cid}.xhtml'
            body = f'<h1>{chap["title"]}</h1>\n{chap["content"]}'
            xhtml = (f'<?xml version="1.0" encoding="utf-8"?>'
                     f'<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">'
                     f'<html xmlns="http://www.w3.org/1999/xhtml"><head><title>{chap["title"]}</title>'
                     f'<link rel="stylesheet" type="text/css" href="../Styles/style.css"/></head>'
                     f'<body>{body}</body></html>')
            chap_files[f'OEBPS/{fname}'] = xhtml.encode('utf-8')
            ncx_pts  += (f'<navPoint id="np{i+1}" playOrder="{i+1}"><navLabel><text>{chap["title"]}</text>'
                         f'</navLabel><content src="{fname}"/></navPoint>\n')
            opf_items += f'<item id="{cid}" href="{fname}" media-type="application/xhtml+xml"/>\n'
            opf_spine += f'<itemref idref="{cid}"/>\n'

        opf = (f'<?xml version="1.0" encoding="utf-8"?>'
               f'<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="bookid" version="2.0">'
               f'<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
               f'<dc:title>{title}</dc:title><dc:creator>{author}</dc:creator>'
               f'<dc:identifier id="bookid">{uid}</dc:identifier><dc:language>en</dc:language>'
               f'</metadata><manifest>'
               f'<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>'
               f'<item id="css" href="Styles/style.css" media-type="text/css"/>'
               f'{opf_items}</manifest><spine toc="ncx">{opf_spine}</spine></package>')
        ncx = (f'<?xml version="1.0" encoding="utf-8"?>'
               f'<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">'
               f'<head><meta name="dtb:uid" content="{uid}"/></head>'
               f'<docTitle><text>{title}</text></docTitle>'
               f'<navMap>{ncx_pts}</navMap></ncx>')
        container = ('<?xml version="1.0"?><container version="1.0" '
                     'xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
                     '<rootfiles><rootfile full-path="OEBPS/content.opf" '
                     'media-type="application/oebps-package+xml"/></rootfiles></container>')
        with zipfile.ZipFile(fp, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('mimetype', 'application/epub+zip', zipfile.ZIP_STORED)
            zf.writestr('META-INF/container.xml', container)
            zf.writestr('OEBPS/content.opf',    opf)
            zf.writestr('OEBPS/toc.ncx',        ncx)
            zf.writestr('OEBPS/Styles/style.css', css)
            for fname, data in chap_files.items():
                zf.writestr(fname, data)

    # ── IMPORT v2 ────────────────────────────────────────────────────────────

    def import_document_v2(self):
        """Open file dialog and return preview chapters — no side effects."""
        if not getattr(self, 'current_project_path', None):
            return {'success': False, 'error': 'No project open.'}

        file_types = ('Supported Files (*.txt;*.md;*.docx;*.html;*.fountain)',
                      'All Files (*.*)')
        result = self._window.create_file_dialog(
            webview.FileDialog.OPEN, file_types=file_types, allow_multiple=False)
        if not result or len(result) == 0:
            return {'success': False, 'cancelled': True}

        filepath = result[0]
        ext = os.path.splitext(filepath)[1].lower()

        try:
            if ext in ('.txt', '.html'):
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    raw = f.read()
                if ext == '.html':
                    import re
                    raw = re.sub(r'<[^>]+>', ' ', raw)
                chapters = self._split_raw_text_into_chapters(raw, 'Imported Chapter')

            elif ext == '.md':
                chapters = self._import_md(filepath)

            elif ext == '.fountain':
                chapters = self._import_fountain(filepath)

            elif ext == '.docx':
                try:
                    import docx
                    doc   = docx.Document(filepath)
                    raw   = '\n'.join(p.text for p in doc.paragraphs)
                    chapters = self._split_raw_text_into_chapters(raw, 'Imported Chapter')
                except ImportError:
                    return {'success': False, 'error': 'python-docx required for .docx import.'}
            else:
                return {'success': False, 'error': f'Unsupported format: {ext}'}

            if not chapters:
                return {'success': False, 'error': 'No content found in file.'}

            # Return preview: titles + first 120 chars of content
            preview = [{'title': c['title'],
                        'preview': (c.get('content','')[:120] + '…') if len(c.get('content','')) > 120 else c.get('content',''),
                        'content': c['content']}
                       for c in chapters]
            return {'success': True, 'chapters': preview,
                    'filename': os.path.basename(filepath)}

        except Exception as e:
            import traceback; traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def _import_md(self, filepath):
        """Split a Markdown file into chapters on ## headings."""
        import re
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        parts = re.split(r'^##\s+', text, flags=re.MULTILINE)
        chapters = []
        for i, part in enumerate(parts):
            if not part.strip(): continue
            lines = part.split('\n', 1)
            ctitle = lines[0].strip() or f'Chapter {i}'
            body   = lines[1].strip() if len(lines) > 1 else ''
            # Convert basic MD to paragraphs
            body = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', body)
            body = re.sub(r'\*(.*?)\*',     r'<em>\1</em>', body)
            paras = [f'<p>{p.strip()}</p>' for p in body.split('\n\n') if p.strip()]
            chapters.append({'title': ctitle, 'content': '\n'.join(paras)})
        return chapters

    def _import_fountain(self, filepath):
        """Parse Fountain screenplay — group scene headings as chapter breaks."""
        import re
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        chapters = []
        current_title = 'Scene 1'
        current_lines = []

        def flush():
            if current_lines:
                body = '\n'.join(current_lines).strip()
                paras = [f'<p>{l.strip()}</p>' for l in body.split('\n\n') if l.strip()]
                chapters.append({'title': current_title, 'content': '\n'.join(paras)})

        scene_re = re.compile(r'^(INT\.|EXT\.|INT/EXT\.|I/E\.|\.)(.+)', re.IGNORECASE)
        for line in lines:
            stripped = line.rstrip()
            if scene_re.match(stripped.lstrip()):
                flush()
                current_title = stripped.strip().lstrip('.')
                current_lines = []
            else:
                current_lines.append(stripped)
        flush()
        return chapters if chapters else [{'title': 'Imported Screenplay',
                                            'content': '<p>' + ' '.join(l.strip() for l in lines) + '</p>'}]

    # ── PROJECT DASHBOARD ─────────────────────────────────────────────────────

    def scan_projects_folder(self, folder_path):
        """Scan a directory for .inkwell project folders. Returns card data list."""
        import glob
        if not folder_path or not os.path.isdir(folder_path):
            return {'success': False, 'error': 'Invalid folder'}
        cards = []
        for entry in os.scandir(folder_path):
            if not entry.is_dir() or not entry.name.endswith('.inkwell'):
                continue
            # Find the meta JSON (any .json in the folder root)
            meta_files = glob.glob(os.path.join(entry.path, '*.json'))
            meta = None
            for mf in meta_files:
                if 'checkpoint' not in mf.lower():
                    try:
                        with open(mf, 'r', encoding='utf-8') as f:
                            meta = json.load(f)
                        break
                    except Exception:
                        continue
            if not meta:
                continue
            # Word count
            wc = 0
            for chap in meta.get('chapters', []):
                raw = re.sub(r'<[^>]+>', ' ', chap.get('content', ''))
                wc += len(raw.split())
            mod_ts = os.path.getmtime(entry.path)
            from datetime import date as _date
            mod_date = datetime.fromtimestamp(mod_ts).strftime('%b %d, %Y')
            cards.append({
                'path':      entry.path,
                'meta_file': meta_files[0] if meta_files else '',
                'title':     meta.get('title', entry.name.replace('.inkwell', '')),
                'chapters':  len(meta.get('chapters', [])),
                'words':     wc,
                'modified':  mod_date,
                'status':    meta.get('status', ''),
            })
        cards.sort(key=lambda c: c['modified'], reverse=True)
        return {'success': True, 'cards': cards}

    def open_project_from_dashboard(self, meta_file_path):
        """Load a project from the dashboard — same as open_project."""
        return self.open_project(meta_file_path)

    def open_project_dashboard(self):
        """Open (or focus) the multi-project dashboard satellite."""
        import webview
        for w in webview.windows:
            if w.title == 'Project Dashboard':
                w.restore(); return True

        dash_html = self._inject_holo_mode(self._load_html('project_dashboard.html'))
        dash_win  = webview.create_window(
            'Project Dashboard', html=dash_html, js_api=self,
            width=900, height=620, frameless=True, easy_drag=False,
            resizable=True, background_color='#1A1D23')

        def on_loaded():
            saved_theme = self.get_app_theme()
            dash_win.evaluate_js(f"if(window.syncTheme){{window.syncTheme('{saved_theme}');}}")

        dash_win.events.loaded += on_loaded
        return True

    # Ensure the name is exactly 'analyze_prose'
    def analyze_prose(self, text):
        try:
            # We pass 'self' so the spoke can read your API Keys and URL settings
            return get_llm_suggestions(text, self)
            
        except Exception as e:
            print(f"Grammar Engine Error: {e}")
            return {"success": False, "error": str(e)}
        
    # --- UNIVERSAL SATELLITE WINDOW CONTROLS ---
    def minimize_satellite(self, window_title):
        import webview
        for w in webview.windows:
            if w.title == window_title:
                w.minimize()
                return

    def maximize_satellite(self, window_title):
        import webview
        for w in webview.windows:
            if w.title == window_title:
                w.toggle_fullscreen()
                return

    def close_satellite_window(self, window_title):
        import webview
        for w in webview.windows:
            if w.title == window_title:
                w.destroy()
                return
    
    def resize_satellite(self, window_title, width, height):
        """Pure Python/JS fallback for resizing frameless windows"""
        import webview
        for w in webview.windows:
            if w.title == window_title:
                # Resize while preventing the window from crushing into nothing
                w.resize(max(400, int(width)), max(300, int(height)))
                return

    def open_character_bible(self):
        html = self._inject_holo_mode(self._load_html('character_bible.html'))
        import webview
        # CRITICAL ADDITION: frameless=True removes the ugly Windows/Mac border
        vault_win = webview.create_window(
            'Character Vault',
            html=html,
            width=900,
            height=700, 
            js_api=self, 
            frameless=True,
            easy_drag=False,
            resizable=True,
            background_color='#1A1D23'
        )
        def on_vault_loaded():
            saved_theme = self.get_app_theme()
            vault_win.evaluate_js(f"if(window.syncTheme) {{ window.syncTheme('{saved_theme}'); }}")
            pending = getattr(self, '_pending_vault_focus', None)
            if pending:
                self._pending_vault_focus = None
                import json as _json
                vault_win.evaluate_js(f"setTimeout(function(){{if(window.focusOnCharacter)window.focusOnCharacter({_json.dumps(pending)});}},400);")

        vault_win.events.loaded += on_vault_loaded
        return True

    def open_world_notes(self):
        """Opens the satellite window for World Notes (Codex 2.0)"""
        import webview
        
        for w in webview.windows:
            if w.title == 'World Notes':
                w.restore()
                return

        notes_html = self._inject_holo_mode(self._load_html('world_notes.html'))

        notes_win = webview.create_window(
            'World Notes',
            html=notes_html,
            js_api=self,
            width=900,
            height=600,
            frameless=True,
            easy_drag=False,
            resizable=True,
            background_color='#1A1D23'
        )

        def on_notes_loaded():
            saved_theme = self.get_app_theme()
            notes_win.evaluate_js(f"if(window.syncTheme) {{ window.syncTheme('{saved_theme}'); }}")

        notes_win.events.loaded += on_notes_loaded

    def open_stats_window(self):
        """Opens (or focuses) the Writing Stats satellite window."""
        import webview
        for w in webview.windows:
            if w.title == 'Writing Stats':
                w.restore()
                return True

        stats_html = self._inject_holo_mode(self._load_html('stats_window.html'))
        stats_win = webview.create_window(
            'Writing Stats',
            html=stats_html,
            js_api=self,
            width=720,
            height=540,
            frameless=True,
            easy_drag=False,
            resizable=True,
            background_color='#1A1D23'
        )

        def on_stats_loaded():
            saved_theme = self.get_app_theme()
            stats_win.evaluate_js(f"if(window.syncTheme) {{ window.syncTheme('{saved_theme}'); }}")

        stats_win.events.loaded += on_stats_loaded
        return True

    # ── RELATIONSHIP WEB ────────────────────────────────────────────────────

    def get_relationship_web(self):
        """Return character list, world note entries, and saved web data for the
        relationship web satellite window."""
        if not getattr(self, 'current_project_path', None):
            return {'characters': [], 'worldNotes': [], 'web': {'nodes': [], 'edges': []}}

        # ── Characters ────────────────────────────────────────────────────
        char_path = os.path.join(self.current_project_path, 'characters.json')
        characters = []
        if os.path.exists(char_path):
            try:
                with open(char_path, 'r', encoding='utf-8') as f:
                    raw = json.load(f)
                characters = [
                    {'id': c.get('id', ''), 'name': c.get('name', 'Unknown'),
                     'role': c.get('role', ''), 'avatar': c.get('avatar', '')}
                    for c in raw if c.get('id')
                ]
            except Exception as e:
                print(f"[relationship_web] char load error: {e}")

        # ── World notes ────────────────────────────────────────────────────
        # Format: { "Locations": [{id, title, content}], "Factions": [...], ... }
        notes_path = os.path.join(self.current_project_path, 'world_notes.json')
        world_notes = []
        if os.path.exists(notes_path):
            try:
                with open(notes_path, 'r', encoding='utf-8') as f:
                    raw_notes = json.load(f)

                def _strip_html(h):
                    return re.sub(r'<[^>]+>', ' ', str(h or '')).strip()[:120]

                def _process_note(entry, category):
                    if not isinstance(entry, dict):
                        return
                    entry_id = entry.get('id') or entry.get('title', '')
                    if not entry_id:
                        return
                    world_notes.append({
                        'id':       f"wn_{entry_id}",
                        'title':    entry.get('title', 'Untitled'),
                        'category': category,
                        'desc':     _strip_html(entry.get('content', '')),
                        'image':    entry.get('image', ''),
                    })

                if isinstance(raw_notes, dict):
                    for category, entries in raw_notes.items():
                        if not isinstance(entries, list):
                            continue
                        for entry in entries:
                            _process_note(entry, category)
                elif isinstance(raw_notes, list):
                    for entry in raw_notes:
                        _process_note(entry, entry.get('category', 'Lore'))
            except Exception as e:
                print(f"[relationship_web] world notes load error: {e}")

        # ── Saved web layout ───────────────────────────────────────────────
        web_path = os.path.join(self.current_project_path, 'relationship_web.json')
        web_data = {'nodes': [], 'edges': []}
        if os.path.exists(web_path):
            try:
                with open(web_path, 'r', encoding='utf-8') as f:
                    web_data = json.load(f)
            except Exception as e:
                print(f"[relationship_web] web load error: {e}")

        return {'characters': characters, 'worldNotes': world_notes, 'web': web_data}

    def save_relationship_web(self, data):
        """Persist node positions and edges for the relationship web."""
        if not getattr(self, 'current_project_path', None):
            return False
        web_path = os.path.join(self.current_project_path, 'relationship_web.json')
        try:
            with open(web_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2)
            return True
        except Exception as e:
            print(f"[save_relationship_web] {e}")
            return False

    def open_relationship_web(self):
        """Opens (or focuses) the Character Relationship Web satellite window."""
        import webview
        for w in webview.windows:
            if w.title == 'Relationship Web':
                w.restore()
                return True

        html = self._inject_holo_mode(self._load_html('relationship_web.html'))
        win  = webview.create_window(
            'Relationship Web',
            html=html,
            js_api=self,
            width=900,
            height=680,
            frameless=True,
            easy_drag=False,
            resizable=True,
            background_color='#1A1D23',
        )

        def on_loaded():
            saved_theme = self.get_app_theme()
            win.evaluate_js(f"if(window.syncTheme){{window.syncTheme('{saved_theme}');}}")
            win.evaluate_js("loadData();")
            # If a focus was queued before the window finished loading, apply it now
            pending = getattr(self, '_pending_web_focus', None)
            if pending:
                self._pending_web_focus = None
                import json as _json
                win.evaluate_js(f"setTimeout(function(){{if(window.focusNodeById)window.focusNodeById({_json.dumps(pending)});}},600);")

        win.events.loaded += on_loaded
        return True

    def open_relationship_web_focused(self, node_id):
        """Open the Relationship Web and focus a specific node once loaded."""
        import webview, json as _json
        # If already open, call focusNodeById directly
        for w in webview.windows:
            if w.title == 'Relationship Web':
                w.restore()
                w.evaluate_js(f"if(window.focusNodeById)window.focusNodeById({_json.dumps(node_id)});")
                return True
        # Not open — store pending focus and open
        self._pending_web_focus = node_id
        return self.open_relationship_web()

    def notify_web_character_update(self, char_id, name, avatar):
        """Push a character name/avatar change to the open Relationship Web window."""
        import webview, json as _json
        payload = _json.dumps({'id': char_id, 'label': name, 'avatar': avatar or ''})
        for w in webview.windows:
            if w.title == 'Relationship Web':
                w.evaluate_js(f"if(window.updateNodeFromVault)window.updateNodeFromVault({payload});")
                break

    def open_character_in_vault(self, char_id):
        """Open Character Vault focused on a specific character."""
        import webview, json as _json
        # If vault already open, call focusOnCharacter directly
        for w in webview.windows:
            if w.title == 'Character Vault':
                w.restore()
                w.evaluate_js(f"if(window.focusOnCharacter)window.focusOnCharacter({_json.dumps(char_id)});")
                return True
        # Not open — store pending focus and open
        self._pending_vault_focus = char_id
        return self.open_character_bible()

    # ── TIMELINE ──────────────────────────────────────────────────────

    def get_timeline_data(self):
        """Return all chapters with POV and metadata for the Timeline satellite."""
        import re as _re
        if not getattr(self, 'current_meta_path', None):
            return {'chapters': []}
        try:
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                db = json.load(f)
        except Exception as e:
            print(f"[get_timeline_data] {e}")
            return {'chapters': []}

        chapters = db.get('chapters', [])
        chap_dir = os.path.join(self.current_project_path, 'chapters') \
                   if getattr(self, 'current_project_path', None) else None

        result = []
        for ch in chapters:
            # Word count: read from the per-chapter content file (no inline content in meta)
            wc = 0
            if chap_dir:
                chap_file = os.path.join(chap_dir, f"{ch.get('id','')}.html")
                if os.path.exists(chap_file):
                    try:
                        with open(chap_file, 'r', encoding='utf-8') as cf:
                            raw = cf.read()
                        text = _re.sub(r'<[^>]*>', ' ', raw)
                        text = _re.sub(r'&[a-z#0-9]+;', ' ', text, flags=_re.IGNORECASE)
                        wc = len([w for w in text.split() if w])
                    except Exception:
                        pass
            result.append({
                'id':        ch.get('id', ''),
                'title':     ch.get('title', 'Untitled'),
                'synopsis':  ch.get('synopsis', ''),
                'status':    ch.get('status', 'draft'),
                'mood':      ch.get('mood', ''),
                'pov':       ch.get('pov', ''),
                'wordCount': wc,
            })
        return {'chapters': result}

    def save_timeline_order(self, data):
        """Reorder chapters and update POV assignments from the Timeline.
        data: list of {id, pov} in the new order.
        Notifies the main window to refresh its binder."""
        import webview, json as _json
        if not getattr(self, 'current_meta_path', None):
            return False
        try:
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                db = json.load(f)
        except Exception as e:
            print(f"[save_timeline_order] read error: {e}")
            return False

        chapters = db.get('chapters', [])
        ch_map = {ch['id']: ch for ch in chapters}

        new_order = []
        seen = set()
        for item in (data or []):
            cid = item.get('id')
            ch  = ch_map.get(cid)
            if ch and cid not in seen:
                ch['pov'] = item.get('pov', ch.get('pov', ''))
                new_order.append(ch)
                seen.add(cid)
        # Preserve any chapters not mentioned (safety net)
        for ch in chapters:
            if ch.get('id') not in seen:
                new_order.append(ch)

        db['chapters'] = new_order
        try:
            with open(self.current_meta_path, 'w', encoding='utf-8') as f:
                json.dump(db, f, indent=2)
        except Exception as e:
            print(f"[save_timeline_order] write error: {e}")
            return False

        # Push new order to the main window so its binder re-renders without a reload
        payload = _json.dumps(data)
        for w in webview.windows:
            if w.title == 'Inkwell':
                w.evaluate_js(f"if(window.onTimelineSync)window.onTimelineSync({payload});")
                break
        return True

    def jump_to_chapter(self, chapter_id):
        """Tell the main window to scroll to and activate a chapter by id."""
        import webview, json as _json
        for w in webview.windows:
            if w.title == 'Inkwell':
                w.evaluate_js(f"if(window.jumpToChapterById)window.jumpToChapterById({_json.dumps(chapter_id)});")
                return True
        return False

    def open_timeline(self):
        """Opens (or focuses) the Story Timeline satellite window."""
        import webview
        for w in webview.windows:
            if w.title == 'Timeline':
                w.restore()
                return True

        html = self._inject_holo_mode(self._load_html('timeline.html'))
        win  = webview.create_window(
            'Timeline',
            html=html,
            js_api=self,
            width=1050,
            height=600,
            frameless=True,
            easy_drag=False,
            resizable=True,
            background_color='#1A1D23',
        )

        def on_loaded():
            saved_theme = self.get_app_theme()
            win.evaluate_js(f"if(window.syncTheme){{window.syncTheme('{saved_theme}');}}")
            win.evaluate_js("loadData();")

        win.events.loaded += on_loaded
        return True

    # ── PHASE 9: REVISION / SNAPSHOT TRACKING ──────────────────────────────

    def _snapshots_dir(self):
        """Return (and create) the snapshots directory for the current project."""
        if not getattr(self, 'current_project_path', None):
            return None
        d = os.path.join(self.current_project_path, 'snapshots')
        os.makedirs(d, exist_ok=True)
        return d

    def save_snapshot(self, name):
        """Save a named snapshot of every chapter's current content.

        Returns { success, id, timestamp, error }
        """
        snap_dir = self._snapshots_dir()
        if not snap_dir:
            return {'success': False, 'error': 'No project open'}

        ts_iso  = datetime.now().isoformat(timespec='seconds')
        snap_id = datetime.now().strftime('%Y%m%d_%H%M%S')
        dest    = os.path.join(snap_dir, snap_id)
        try:
            os.makedirs(dest, exist_ok=True)

            # Read current meta for chapter list
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                db = json.load(f)

            chapters_meta = []
            chap_dir = os.path.join(self.current_project_path, 'chapters')
            for chap in db.get('chapters', []):
                cid   = chap['id']
                title = chap.get('title', f'Chapter {cid}')
                src   = os.path.join(chap_dir, f'{cid}.html')
                content = ''
                if os.path.exists(src):
                    with open(src, 'r', encoding='utf-8') as cf:
                        content = cf.read()
                with open(os.path.join(dest, f'{cid}.html'), 'w', encoding='utf-8') as df:
                    df.write(content)
                chapters_meta.append({'id': cid, 'title': title})

            meta = {
                'id':        snap_id,
                'name':      (name or '').strip() or f'Snapshot {snap_id}',
                'timestamp': ts_iso,
                'chapters':  chapters_meta,
            }
            with open(os.path.join(dest, 'meta.json'), 'w', encoding='utf-8') as mf:
                json.dump(meta, mf, indent=2)

            return {'success': True, 'id': snap_id, 'timestamp': ts_iso}
        except Exception as e:
            print(f'[save_snapshot] {e}')
            return {'success': False, 'error': str(e)}

    def list_snapshots(self):
        """Return all snapshots for the current project, newest first."""
        snap_dir = self._snapshots_dir()
        if not snap_dir:
            return []
        results = []
        try:
            for entry in sorted(os.listdir(snap_dir), reverse=True):
                meta_path = os.path.join(snap_dir, entry, 'meta.json')
                if os.path.isfile(meta_path):
                    with open(meta_path, 'r', encoding='utf-8') as f:
                        results.append(json.load(f))
        except Exception as e:
            print(f'[list_snapshots] {e}')
        return results

    def delete_snapshot(self, snapshot_id):
        """Permanently delete a snapshot folder."""
        snap_dir = self._snapshots_dir()
        if not snap_dir:
            return {'success': False}
        target = os.path.join(snap_dir, snapshot_id)
        try:
            if os.path.isdir(target):
                shutil.rmtree(target)
            return {'success': True}
        except Exception as e:
            return {'success': False, 'error': str(e)}

    def get_snapshot_diff(self, snapshot_id, chapter_id):
        """Compute a word-level diff between the snapshot version and the current
        saved version of one chapter.

        Returns:
            { success, spans: [{ type: 'equal'|'delete'|'insert', text }] }
        """
        import difflib, html as _html
        snap_dir = self._snapshots_dir()
        if not snap_dir:
            return {'success': False, 'error': 'No project open'}

        snap_file = os.path.join(snap_dir, snapshot_id, f'{chapter_id}.html')
        curr_file = os.path.join(self.current_project_path, 'chapters', f'{chapter_id}.html')

        def _strip(path):
            if not os.path.exists(path):
                return ''
            with open(path, 'r', encoding='utf-8') as f:
                raw = f.read()
            import re as _re
            text = _re.sub(r'<[^>]+>', ' ', raw)
            text = _html.unescape(text)
            text = _re.sub(r'\s+', ' ', text).strip()
            return text

        old_text = _strip(snap_file)
        new_text = _strip(curr_file)

        import re as _re
        _TOK = _re.compile(r'(\s+|\w+|[^\w\s])')
        old_toks = _TOK.findall(old_text)
        new_toks = _TOK.findall(new_text)

        sm = difflib.SequenceMatcher(None, old_toks, new_toks, autojunk=False)
        spans = []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                spans.append({'type': 'equal',  'text': ''.join(old_toks[i1:i2])})
            elif tag == 'delete':
                spans.append({'type': 'delete', 'text': ''.join(old_toks[i1:i2])})
            elif tag == 'insert':
                spans.append({'type': 'insert', 'text': ''.join(new_toks[j1:j2])})
            elif tag == 'replace':
                spans.append({'type': 'delete', 'text': ''.join(old_toks[i1:i2])})
                spans.append({'type': 'insert', 'text': ''.join(new_toks[j1:j2])})

        # Merge consecutive equal spans
        merged = []
        for s in spans:
            if merged and merged[-1]['type'] == s['type'] == 'equal':
                merged[-1]['text'] += s['text']
            else:
                merged.append(s)

        return {'success': True, 'spans': merged}

    def open_snapshots_window(self):
        """Opens (or focuses) the Revision Snapshots satellite window."""
        import webview
        for w in webview.windows:
            if w.title == 'Snapshots':
                w.restore()
                return True

        html = self._inject_holo_mode(self._load_html('snapshots.html'))
        win  = webview.create_window(
            'Snapshots',
            html=html,
            js_api=self,
            width=960,
            height=640,
            frameless=True,
            easy_drag=False,
            resizable=True,
            background_color='#1A1D23',
        )

        def on_loaded():
            saved_theme = self.get_app_theme()
            win.evaluate_js(f"if(window.syncTheme){{window.syncTheme('{saved_theme}');}}")
            win.evaluate_js("loadData();")

        win.events.loaded += on_loaded
        return True

    # ── PHASE 9: WORD FREQUENCY HEATMAP ────────────────────────────────────

    def get_word_frequency_data(self, threshold=None):
        """Scan every chapter in the current project, count word frequencies, and
        return the top overused content words.

        Returns:
            { words: [{word, count}], threshold, total_words }
        """
        import re as _re, html as _html
        if not getattr(self, 'current_meta_path', None):
            return {'words': [], 'threshold': 0, 'total_words': 0}

        _STOP = {
            'the','a','an','and','or','but','in','on','at','to','for','of','with',
            'by','from','is','are','was','were','be','been','being','have','has',
            'had','do','does','did','will','would','could','should','may','might',
            'shall','can','that','this','these','those','it','its','he','she','they',
            'we','i','you','him','her','them','us','my','your','his','their','our',
            'not','no','if','as','so','than','then','when','where','which','who',
            'what','how','all','any','some','each','every','there','here','just',
            'also','up','out','about','into','over','after','before','between',
            'through','during','said','say','says','like','more','one','two','three',
            'new','now','also','very','too','even','back','still','only','much',
            'such','own','other','another','both','same','first','last','well',
            'down','way','s','t','re','ve','d','m','ll','don','didn','doesn',
            'isn','wasn','weren','won','hadn','haven','hasn','couldn','shouldn',
            'wouldn','cant','got','get','go','come','came','make','made','take',
            'took','know','knew','see','saw','look','looked','think','thought',
            'want','wanted','need','needed','feel','felt','tell','told',
        }

        try:
            with open(self.current_meta_path, 'r', encoding='utf-8') as f:
                db = json.load(f)
        except Exception:
            return {'words': [], 'threshold': 0, 'total_words': 0}

        chap_dir = os.path.join(self.current_project_path, 'chapters')
        freq = {}
        total = 0

        for chap in db.get('chapters', []):
            cfile = os.path.join(chap_dir, f"{chap['id']}.html")
            if not os.path.exists(cfile):
                continue
            with open(cfile, 'r', encoding='utf-8') as cf:
                raw = cf.read()
            text = _re.sub(r'<[^>]+>', ' ', raw)
            text = _html.unescape(text).lower()
            words = _re.findall(r"[a-z']{3,}", text)
            for w in words:
                w = w.strip("'")
                if w and w not in _STOP and len(w) >= 3:
                    freq[w] = freq.get(w, 0) + 1
                    total += 1

        if not freq:
            return {'words': [], 'threshold': 0, 'total_words': total}

        if threshold is None:
            auto = max(5, int(total * 0.002))
        else:
            try:
                auto = max(1, int(threshold))
            except (TypeError, ValueError):
                auto = 5

        flagged = sorted(
            [{'word': w, 'count': c} for w, c in freq.items() if c >= auto],
            key=lambda x: x['count'],
            reverse=True,
        )[:120]

        return {'words': flagged, 'threshold': auto, 'total_words': total}

    # ── UPDATE PIPELINE ───────────────────────────────────────────────────────

    def get_app_version(self):
        """Return the running application version string."""
        return APP_VERSION

    def get_update_check_url(self):
        """Return the configured update manifest URL (empty string = disabled)."""
        return getattr(self, 'update_check_url', '')

    def save_update_check_url(self, url):
        """Persist a new update manifest URL and return it."""
        self.update_check_url = (url or '').strip()
        self._save_app_config()
        return self.update_check_url

    def _version_tuple(self, v):
        """Convert 'A.B.C' → (A, B, C) for numeric comparison."""
        try:
            return tuple(int(x) for x in str(v).strip().split('.'))
        except Exception:
            return (0, 0, 0)

    def check_for_updates(self):
        """
        Fetch the remote update manifest and compare against APP_VERSION.

        Returns a dict:
          has_update      – bool
          current_version – str
          latest_version  – str  (same as current if no update / on error)
          download_url    – str
          release_notes   – str
          error           – str  (only present when something went wrong)

        The manifest JSON at update_check_url should look like:
          {
            "version":       "1.0.2",
            "download_url":  "https://example.com/Inkwell-Setup-1.0.2.exe",
            "release_notes": "Bug fixes and new features."
          }
        """
        import urllib.request, json as _json

        url = getattr(self, 'update_check_url', '').strip()
        if not url:
            return {
                'has_update': False,
                'current_version': APP_VERSION,
                'latest_version':  APP_VERSION,
                'download_url':    '',
                'release_notes':   '',
                'error':           'no_url',
            }

        try:
            req = urllib.request.Request(
                url,
                headers={'User-Agent': f'Inkwell/{APP_VERSION}'},
            )
            with urllib.request.urlopen(req, timeout=10) as resp:
                data = _json.loads(resp.read().decode('utf-8'))

            latest       = str(data.get('version', '') or '').strip()
            download_url = str(data.get('download_url', '') or '').strip()
            notes        = str(data.get('release_notes', '') or '').strip()

            if not latest:
                raise ValueError('manifest missing "version" field')

            has_update = self._version_tuple(latest) > self._version_tuple(APP_VERSION)

            return {
                'has_update':      has_update,
                'current_version': APP_VERSION,
                'latest_version':  latest,
                'download_url':    download_url,
                'release_notes':   notes,
            }

        except Exception as exc:
            return {
                'has_update':      False,
                'current_version': APP_VERSION,
                'latest_version':  APP_VERSION,
                'download_url':    '',
                'release_notes':   '',
                'error':           str(exc),
            }

    def download_and_install_update(self, download_url):
        """
        Download the installer to a temp file in a background thread,
        stream progress back to the UI via JS callbacks, then launch
        the installer and exit the app.

        JS callbacks used (all optional / guarded with &&):
          window._onUpdateDownloadProgress(pct)  – 0-100
          window._onUpdateReady()                 – installer launched, app about to exit
          window._onUpdateError(msg)              – something went wrong
        """
        import urllib.request, tempfile, threading, time as _time

        if not download_url:
            return False

        win = self._window

        def _push(js):
            try:
                if win:
                    win.evaluate_js(js)
            except Exception:
                pass

        def _worker():
            tmp_path = None
            try:
                # Sanitise URL
                url = download_url.strip()

                # Stream download, report progress
                req = urllib.request.Request(
                    url,
                    headers={'User-Agent': f'Inkwell/{APP_VERSION}'},
                )
                suffix = '.exe' if url.lower().endswith('.exe') else '.tmp'
                tmp_fd, tmp_path = tempfile.mkstemp(suffix=suffix)

                with urllib.request.urlopen(req, timeout=120) as resp, \
                     os.fdopen(tmp_fd, 'wb') as fp:
                    total = int(resp.headers.get('Content-Length') or 0)
                    downloaded = 0
                    chunk = 65536
                    while True:
                        buf = resp.read(chunk)
                        if not buf:
                            break
                        fp.write(buf)
                        downloaded += len(buf)
                        if total > 0:
                            pct = min(99, int(downloaded / total * 100))
                            _push(f'window._onUpdateDownloadProgress && window._onUpdateDownloadProgress({pct})')

                # Launch installer
                flags = subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                subprocess.Popen([tmp_path], creationflags=flags)

                _push('window._onUpdateReady && window._onUpdateReady()')
                _time.sleep(1.8)
                sys.exit(0)

            except Exception as exc:
                err_js = json.dumps(str(exc))
                _push(f'window._onUpdateError && window._onUpdateError({err_js})')
                # Clean up on failure
                if tmp_path and os.path.exists(tmp_path):
                    try:
                        os.remove(tmp_path)
                    except Exception:
                        pass

        threading.Thread(target=_worker, daemon=True).start()
        return True

    def open_url_in_browser(self, url):
        """Open a URL in the system default browser."""
        import webbrowser
        if url:
            webbrowser.open(url)

        return {'words': flagged, 'threshold': auto, 'total_words': total}
